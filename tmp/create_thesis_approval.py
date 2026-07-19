from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/ahmed/Desktop/Thesis_Approval_Certificate.docx"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.top_margin = Inches(1.75)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(1.35)
sec.right_margin = Inches(0.90)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def borderless(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)

def add_run(p, text, bold=False):
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return r

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(13)
add_run(title, "THESIS APPROVAL CERTIFICATE", True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
add_run(p, 'The thesis study of Chemistry Doctorate Program (English) student Orhan YILDIRIM\n'
           'with student number 1234567 titled “Polymerization Of Monoisocyanates By Free\n'
           'Radical Initiation With A Novel Room Temperature Initiator” has been approved by\n'
           'the Jury and has been accepted as a Doctor of Philosophy (PhD) Thesis.')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(28)
add_run(p, "Thesis Defense Date: ", True)
add_run(p, "22/12/2026")

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.autofit = False
widths = [Inches(0.28), Inches(4.27), Inches(1.05)]
for i, w in enumerate(widths):
    table.columns[i].width = w
    table.cell(0, i).width = w
borderless(table)

headers = ["", "Jury Members", "Signature"]
for i, text in enumerate(headers):
    cell = table.cell(0, i)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    if i == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bool(text))

members = [
    ("1)", "Prof. Dr. Savaş YAVUZ", "Chair"),
    ("2)", "Prof. Dr. Ali USANMAZ", "Supervisor"),
    ("3)", "Prof. Dr. Güngör AKGÜNDÜZ", "Member"),
    ("4)", "Assoc. Prof. Dr. Steve BLACK", "Member"),
    ("5)", "Asst. Prof. Dr. Serpil ALSOY", "Member"),
]
for num, name, role in members:
    row = table.add_row()
    for i, w in enumerate(widths):
        row.cells[i].width = w
        set_cell_margins(row.cells[i], top=65, bottom=65)
        row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_run(row.cells[0].paragraphs[0], num)
    add_run(row.cells[1].paragraphs[0], name, True)
    add_run(row.cells[1].add_paragraph(), role)
    sig = row.cells[2].paragraphs[0]
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sig, "………………")

spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(11)

approvals = doc.add_table(rows=0, cols=2)
approvals.alignment = WD_TABLE_ALIGNMENT.LEFT
approvals.autofit = False
aw = [Inches(4.55), Inches(1.05)]
borderless(approvals)

approval_data = [
    ("Prof. Dr. Ali USANMAZ", ["Supervisor"]),
    ("Prof. Dr. Namık Kemal URAS", ["Head of the Chemistry PhD Program", "Head of the Natural Sciences Department"]),
    ("Prof. Dr. Osman YILMAZ", ["Director", "Institute of Graduate Studies and Research", "of Cyprus International University"]),
]
for name, lines in approval_data:
    row = approvals.add_row()
    for i, w in enumerate(aw):
        row.cells[i].width = w
        set_cell_margins(row.cells[i], top=80, bottom=80)
        row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_run(row.cells[0].paragraphs[0], name, True)
    for line in lines:
        add_run(row.cells[0].add_paragraph(), line)
    sig = row.cells[1].paragraphs[0]
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sig, "………………")

doc.save(OUT)
print(OUT)
