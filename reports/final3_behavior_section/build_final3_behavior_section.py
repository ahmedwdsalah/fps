#!/usr/bin/env python3
from __future__ import annotations

import json
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "results" / "final3_behavior_loop"
OUT = ROOT / "reports" / "final3_behavior_section"
DOCX = OUT / "final_three_behavior_section.docx"
MD = OUT / "final_three_behavior_section.md"
FIGS = OUT / "figures"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15 if level == 1 else 12)
    r.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(89, 89, 89)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "1F4E79")
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.9) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)


def copy_figures() -> dict[str, Path]:
    FIGS.mkdir(parents=True, exist_ok=True)
    wanted = {
        "winner": "baseline_winner_share.png",
        "domain": "domain_winner_patterns.png",
        "success": "success_top_indicators.png",
        "failure": "failure_pairs_regret.png",
        "interaction": "size_feature_interaction_patterns.png",
    }
    out = {}
    for key, name in wanted.items():
        src = SOURCE / "figures" / name
        dst = FIGS / name
        if src.exists():
            shutil.copy2(src, dst)
            out[key] = dst
    return out


def build_markdown(summary: dict, domain: pd.DataFrame, failures: pd.DataFrame, worst: pd.DataFrame) -> str:
    return f"""# Final-Three Runtime Behaviour Analysis

This section extends the portfolio reduction study after the candidate set has been reduced to introsort, heapsort, and timsort. The purpose is not to introduce a new algorithm space, but to explain how the retained portfolio behaves after labelling and prediction are restricted to the final three classes.

The analysis uses {summary['rows']:,} labelled arrays from the existing training dataset. For each array, the fastest algorithm is recomputed from the measured runtimes of introsort, heapsort, and timsort. The virtual best solver total runtime is {summary['vbs_total_s']:.6f} seconds, while the single best solver total runtime is {summary['sbs_total_s']:.6f} seconds using {summary['sbs_algorithm']}. This gives a remaining SBS-to-VBS gap of {summary['vbs_sbs_gap_pct']:.2f} percent before per-instance selection is considered.

The winner distribution confirms that the final learning problem is highly structured rather than uniform. Timsort is fastest on 79.51 percent of the rows, heapsort on 15.40 percent, and introsort on 5.09 percent. This imbalance explains why classification accuracy alone is not sufficient: a selector may appear strong by following the largest class, while still producing expensive regret in smaller or boundary regions.

Domain behaviour also shows that the same three-algorithm portfolio does not behave identically across sources. Crypto, Stock, and Earthquake are dominated by timsort, while the smaller Weather subset is more mixed and has lower prediction success. Weather therefore behaves as a harder transfer region, not because the number of elements alone is larger or smaller, but because the structural indicators and timing margins differ.

The strongest correct predictions occur where the structural evidence is clear. High sortedness, low inversion ratio, and stable run structure produce regions where timsort dominates above 99 percent and the model succeeds almost perfectly. Heapsort correctness is associated with high duplicate and frequency concentration, especially high top-5 frequency ratio and high duplicate ratio. Introsort appears more often in larger and less cleanly ordered regions, but because it is the smallest winner class, its boundary is less stable.

Failure analysis separates harmless mistakes from meaningful mistakes. Out of all predictions, 36,332 are failures, but 19,806 of them have regret at or below one microsecond. These are boundary decisions where two algorithms are practically tied. The more important errors are high-regret failures, especially timsort predicted as introsort, heapsort predicted as introsort, and timsort predicted as heapsort. These cases have a different structural profile: larger median size, weaker longest-run signal, higher inversion ratio, and much lower duplicate/frequency concentration than ordinary failures.

This result supports the main evaluation logic of the thesis. The selector should not be judged only by whether the predicted label matches the fastest label. It should be judged by the runtime penalty of the mistake. Correctness is easiest where structural indicators are strong; failure becomes important when the model crosses a boundary between algorithms whose measured runtimes are no longer near-tied.
"""


def build_docx() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    figs = copy_figures()
    tables = SOURCE / "tables"
    summary = json.loads((tables / "baseline_summary.json").read_text())
    winner = pd.read_csv(tables / "baseline_winner_share.csv")
    domain = pd.read_csv(tables / "domain_winner_success_patterns.csv")
    failures = pd.read_csv(tables / "failure_pairs.csv")
    worst = pd.read_csv(tables / "worst_failure_pair_anatomy.csv")
    shifts = pd.read_csv(tables / "worst_failure_feature_shift.csv")
    class_patterns = pd.read_csv(tables / "class_specific_success_patterns.csv")

    MD.write_text(build_markdown(summary, domain, failures, worst), encoding="utf-8")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    add_heading(doc, "Final-Three Runtime Behaviour Analysis")
    add_para(doc, "After the wider portfolio study had reduced the candidate space, the retained algorithms were examined as a separate runtime behaviour problem. The purpose of this analysis was to explain the structure of the final three-class selection task rather than only report aggregate model accuracy. The study therefore recomputes the fastest label from the measured runtimes of introsort, heapsort, and timsort, then relates these labels to structural features, prediction success, failure type, and regret.")
    add_para(doc, f"The analysis uses {summary['rows']:,} labelled arrays from the existing training dataset. The virtual best solver total runtime is {summary['vbs_total_s']:.6f} seconds, while the single best solver total runtime is {summary['sbs_total_s']:.6f} seconds using {summary['sbs_algorithm']}. The remaining SBS-to-VBS gap is {summary['vbs_sbs_gap_pct']:.2f} percent, which means that the final three algorithms still contain meaningful room for per-instance selection.")

    add_heading(doc, "Winner Distribution Inside the Final Portfolio", level=2)
    rows = [[r.algorithm, f"{int(r['count']):,}", f"{r.pct:.2f}%"] for _, r in winner.iterrows()]
    add_table(doc, ["Algorithm", "Fastest cases", "Share"], rows)
    add_para(doc, "The winner distribution is not balanced. Timsort is the dominant fastest algorithm, but heapsort and introsort still form non-trivial regions of the empirical space. This confirms that the final task is not a simple fixed-algorithm problem: the majority class is strong, but using only the majority class would leave measurable runtime regret in the regions where the other algorithms are faster.")
    add_figure(doc, figs.get("winner", Path()), "Figure X. Fastest-algorithm share inside the final three-algorithm portfolio.")

    add_heading(doc, "Domain-Level Behaviour", level=2)
    drows = []
    for _, r in domain.iterrows():
        drows.append([r.domain, f"{int(r.n):,}", r.dominant, f"{r.dominant_pct:.1f}%", f"{r.success_pct:.2f}%", f"{r.mean_regret_us:.2f}"])
    add_table(doc, ["Domain", "Rows", "Dominant", "Dominant share", "Prediction success", "Mean regret (us)"], drows)
    add_para(doc, "The final portfolio behaves differently across the data sources. Crypto, Stock, and Earthquake are dominated by timsort, while Weather is more mixed and has lower success. This does not mean that element count alone determines regret. The harder cases are created by the combination of source structure, timing margin, and the algorithm selected by the classifier.")
    add_figure(doc, figs.get("domain", Path()), "Figure X. Winner distribution by domain for the final three algorithms.")

    add_heading(doc, "Structural Conditions Behind Correct Prediction", level=2)
    add_para(doc, "Correct prediction is strongest when the measured runtime winner is supported by clear structural evidence. High sortedness, low inversion ratio, and consistent run structure create regions where timsort dominates almost completely. In these regions the model is not merely learning a label frequency; it is detecting the same structural pattern that makes the runtime measurement favour timsort.")
    top_success = class_patterns.groupby("true_class").head(4)
    rows = []
    for _, r in top_success.iterrows():
        rows.append([r.true_class, r.feature, str(r.bucket), f"{r.success_pct:.1f}%", f"{r.median_margin_us:.2f}"])
    add_table(doc, ["True class", "Feature", "Region", "Success", "Median margin (us)"], rows)
    add_para(doc, "Heapsort success is linked with high duplicate and frequency concentration, especially top-5 frequency and duplicate-ratio regions. Introsort appears in larger and less cleanly ordered cases, but because it is the smallest true winner class, its decision boundary is less stable. This explains why per-class behaviour is more informative than one aggregate accuracy number.")
    add_figure(doc, figs.get("success", Path()), "Figure X. Strongest structural regions behind successful predictions.")

    add_heading(doc, "Failure Structure and Regret", level=2)
    frows = []
    for _, r in failures.iterrows():
        frows.append([r.pair, f"{int(r.n):,}", f"{r.mean_regret_us:.2f}", f"{r.median_regret_us:.2f}", f"{r.p95_regret_us:.2f}"])
    add_table(doc, ["Failure pair", "Cases", "Mean regret (us)", "Median regret (us)", "P95 regret (us)"], frows)
    add_para(doc, "The failure analysis separates label error from runtime damage. There were 36,332 failed predictions, but 19,806 of them had regret at or below one microsecond. These cases are near-boundary decisions where two algorithms are practically tied. The important part of the error distribution is therefore the high-regret subset, not every incorrect label equally.")
    add_figure(doc, figs.get("failure", Path()), "Figure X. Failed prediction pairs ordered by mean regret.")

    add_heading(doc, "Worst-Failure Anatomy", level=2)
    wrows = []
    for _, r in worst.iterrows():
        wrows.append([r.pair, f"{int(r.n):,}", f"{r.mean_regret_us:.2f}", f"{r.median_n:.0f}", f"{r.median_margin_us:.2f}"])
    add_table(doc, ["Failure pair", "Cases", "Mean regret (us)", "Median n", "Median margin (us)"], wrows)
    add_para(doc, "The worst 500 failures have a distinct profile. Their median array size is 13,320.5 compared with 2,500 for all failures and 501 for the full dataset. They also show weaker longest-run structure, higher inversion ratio, and much lower duplicate/frequency concentration than ordinary failures. This indicates that the highest regret cases are not random noise; they are concentrated in structurally difficult regions where the class boundary is expensive when crossed incorrectly.")
    srows = []
    for _, r in shifts.iterrows():
        srows.append([r.feature, f"{r.worst_median:.4g}", f"{r.all_failure_median:.4g}", f"{r.all_data_median:.4g}"])
    add_table(doc, ["Feature", "Worst failures", "All failures", "All data"], srows)

    add_heading(doc, "Interpretation", level=2)
    add_para(doc, "This final-three behaviour study supports the same conclusion as the wider algorithm portfolio reduction, but at a more detailed level. The retained algorithms are not interchangeable labels. Timsort captures ordered and run-adaptive structure, heapsort captures duplicate-heavy and frequency-concentrated regions, and introsort appears in a smaller but still meaningful part of the structural space. The selector succeeds when those structural signals are clear and fails mainly near boundaries or in larger high-regret regions where the wrong class carries measurable runtime cost.")
    add_para(doc, "The analysis also explains why the thesis evaluates runtime value rather than classification accuracy alone. A failed label can be harmless when two algorithms are tied, while a smaller number of errors can dominate regret if the selected algorithm is far from the oracle runtime. The practical question is therefore not only whether the classifier predicts the fastest label, but whether its mistakes preserve most of the SBS-to-VBS runtime gap.")

    doc.save(DOCX)


if __name__ == "__main__":
    build_docx()
    print(DOCX)
