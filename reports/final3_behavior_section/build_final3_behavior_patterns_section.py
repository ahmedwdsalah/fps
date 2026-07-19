#!/usr/bin/env python3
from __future__ import annotations

import json
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "results" / "final3_behavior_loop"
OUT = ROOT / "reports" / "final3_behavior_section"
FIGS = OUT / "patterns_figures"
DOCX = OUT / "final_three_success_failure_patterns_section.docx"
MD = OUT / "final_three_success_failure_patterns_section.md"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: float = 9.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15 if level == 1 else 12.5)
    r.font.color.rgb = RGBColor(31, 78, 121)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], "1F4E79")
        cell_text(t.rows[0].cells[i], h, True, "FFFFFF", 8.8)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=8.8)
    doc.add_paragraph()


def figure(doc: Document, path: Path, caption: str, width: float = 5.9) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(89, 89, 89)


def copy_figures() -> dict[str, Path]:
    FIGS.mkdir(parents=True, exist_ok=True)
    names = {
        "winner": "baseline_winner_share.png",
        "domain": "domain_winner_patterns.png",
        "success": "success_top_indicators.png",
        "failure": "failure_pairs_regret.png",
        "interaction": "size_feature_interaction_patterns.png",
    }
    paths = {}
    for key, name in names.items():
        src = SOURCE / "figures" / name
        dst = FIGS / name
        if src.exists():
            shutil.copy2(src, dst)
            paths[key] = dst
    return paths


def load() -> dict[str, object]:
    base = SOURCE / "tables"
    return {
        "summary": json.loads((base / "baseline_summary.json").read_text()),
        "winner": pd.read_csv(base / "baseline_winner_share.csv"),
        "domain": pd.read_csv(base / "domain_winner_success_patterns.csv"),
        "success": pd.read_csv(base / "class_specific_success_patterns.csv"),
        "regions": pd.read_csv(base / "winner_regions_by_feature_bucket.csv"),
        "failures": pd.read_csv(base / "failure_pairs.csv"),
        "worst": pd.read_csv(base / "worst_failure_pair_anatomy.csv"),
        "shift": pd.read_csv(base / "worst_failure_feature_shift.csv"),
        "interaction": pd.read_csv(base / "size_feature_interaction_patterns.csv"),
    }


def fmt_pct(x: float) -> str:
    return f"{x:.2f}%"


def write_markdown(d: dict[str, object]) -> None:
    summary = d["summary"]
    winner = d["winner"]
    failures = d["failures"]
    tim = float(winner.loc[winner["algorithm"].eq("timsort"), "pct"].iloc[0])
    heap = float(winner.loc[winner["algorithm"].eq("heapsort"), "pct"].iloc[0])
    intro = float(winner.loc[winner["algorithm"].eq("introsort"), "pct"].iloc[0])
    main_fail = failures.iloc[0]
    MD.write_text(
        f"""# Success and Failure Patterns in the Final Three-Algorithm Selector

The final three-algorithm analysis uses {summary['rows']:,} arrays and recomputes the oracle label from measured introsort, heapsort, and timsort runtimes. Timsort is fastest on {tim:.2f}% of cases, heapsort on {heap:.2f}%, and introsort on {intro:.2f}%. The strongest success regions are ordered/run-structured regions for timsort and duplicate/frequency-concentrated regions for heapsort. The dominant failure pair is {main_fail['pair']} with {int(main_fail['n']):,} cases, but the highest regret errors are concentrated in a smaller worst-failure subset with larger size, weaker longest-run structure, higher inversion, and lower duplicate/frequency concentration.
""",
        encoding="utf-8",
    )


def build() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    figs = copy_figures()
    d = load()
    write_markdown(d)
    summary = d["summary"]
    winner = d["winner"]
    domain = d["domain"]
    success = d["success"]
    regions = d["regions"]
    failures = d["failures"]
    worst = d["worst"]
    shift = d["shift"]
    interaction = d["interaction"]

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    heading(doc, "Success and Failure Patterns in the Final Three-Algorithm Selector")
    para(doc, "After the wider algorithm portfolio had been reduced, the remaining question was not only which of introsort, heapsort, and timsort wins most often, but why the selector succeeds in some regions and fails in others. This section therefore treats the final three-algorithm problem as a behavioural analysis of the measured runtime landscape. The fastest label is recomputed from the three measured runtimes, then the prediction result is examined through feature regions, domain composition, failure pairs, and regret magnitude.")
    para(doc, f"The analysis uses {summary['rows']:,} labelled arrays from the archived training dataset. The virtual best solver obtains a total runtime of {summary['vbs_total_s']:.6f} seconds, while the single best solver obtains {summary['sbs_total_s']:.6f} seconds by always choosing {summary['sbs_algorithm']}. The remaining SBS-to-VBS gap is {summary['vbs_sbs_gap_pct']:.2f} percent. This gap is the useful space in which the selector can improve over a fixed algorithm.")

    heading(doc, "Empirical Runtime Landscape", 2)
    rows = [[r.algorithm, f"{int(r['count']):,}", fmt_pct(r.pct)] for _, r in winner.iterrows()]
    table(doc, ["Algorithm", "Fastest cases", "Share"], rows)
    para(doc, "The distribution is strongly skewed toward timsort, but it is not a single-algorithm problem. Heapsort and introsort together account for more than one fifth of the fastest labels. This is the important empirical condition for algorithm selection: the majority class is clear, but there are still measurable regions where a fixed timsort decision would be wrong. The task is therefore to identify when the structure moves an array away from the dominant timsort region.")
    figure(doc, figs.get("winner", Path()), "Figure X. Fastest-algorithm share after reducing the portfolio to introsort, heapsort, and timsort.")

    heading(doc, "Why Correct Predictions Succeed", 2)
    top_t = success[success["true_class"].eq("timsort")].head(4)
    top_h = success[success["true_class"].eq("heapsort")].head(4)
    top_i = success[success["true_class"].eq("introsort")].head(4)
    rows = []
    for label, sub in [("timsort", top_t), ("heapsort", top_h), ("introsort", top_i)]:
        for _, r in sub.iterrows():
            rows.append([label, r.feature, str(r.bucket), f"{r.success_pct:.1f}%", f"{r.median_margin_us:.2f}"])
    table(doc, ["True class", "Dominant feature region", "Observed range", "Success", "Median margin (us)"], rows)
    para(doc, "The strongest timsort successes occur in regions with high adjacent sortedness, low inversion ratio, and stable run structure. This agrees with the mechanism of timsort described in the thesis: timsort is not only comparing values, but first detecting natural ordered runs and then merging them. When local sortedness and run structure are visible in the feature vector, the model has a direct structural signal that corresponds to the algorithm's real advantage.")
    para(doc, "Heapsort successes have a different profile. They are associated with high duplicate and frequency concentration, especially high top-5 frequency and high duplicate-ratio regions. In these regions the array contains repeated values or low-diversity structure, and the measured winner is often not explained by order alone. The model succeeds because the repetition features separate these cases from ordinary partially ordered arrays.")
    para(doc, "Introsort is the hardest successful class to explain because it occupies the smallest share of the empirical label space. Its best regions are larger and less cleanly ordered than the strongest timsort regions, but they do not form as broad or as stable a structural family. This supports the earlier thesis interpretation that the selector is strongest when the algorithmic mechanism leaves a visible feature trace and weaker where the difference depends on lower-level execution effects.")
    figure(doc, figs.get("success", Path()), "Figure X. Structural regions with the strongest successful prediction behaviour.")

    heading(doc, "Size and Structure Work Together", 2)
    top_inter = interaction[interaction["n"] >= 10000].head(8)
    rows = []
    for _, r in top_inter.iterrows():
        rows.append([r.interaction.replace("n_elements x ", ""), str(r.size_range), str(r.feature_range), r.dominant, f"{r.dominant_pct:.1f}%", f"{r.success_pct:.1f}%"])
    table(doc, ["Structure feature", "Size range", "Feature range", "Dominant", "Dominance", "Success"], rows)
    para(doc, "The interaction results show why size cannot be interpreted alone. The same feature can matter differently across size ranges, and the same size range can produce different winners depending on order, inversion, runs, and repetition. This is why the classifier is formulated as a tabular XGBoost model: tree splits can model conditional combinations such as size together with inversion ratio, rather than treating each feature as an isolated linear effect.")
    figure(doc, figs.get("interaction", Path()), "Figure X. Strong size-by-structure regions found in the final three-algorithm analysis.")

    heading(doc, "Why Failed Predictions Fail", 2)
    frows = []
    for _, r in failures.iterrows():
        frows.append([r.pair, f"{int(r.n):,}", f"{r.mean_regret_us:.2f}", f"{r.median_regret_us:.2f}", f"{r.p95_regret_us:.2f}"])
    table(doc, ["Failure pair", "Cases", "Mean regret (us)", "Median regret (us)", "P95 regret (us)"], frows)
    para(doc, "The failure table shows that label error and runtime damage are not the same thing. The most frequent failure pair is timsort_to_heapsort, but its mean regret is smaller than the timsort_to_introsort pair. This means that a failure count alone would give an incomplete interpretation. A selector can make many mistakes in a low-cost boundary region, while fewer mistakes in a high-margin region can produce more practical damage.")
    para(doc, "The low-regret errors are best understood as near-tie decisions. In these cases the measured fastest algorithm and the predicted algorithm are different labels, but their runtimes are almost indistinguishable. Such rows reduce classification accuracy but do not strongly reduce runtime value. This directly supports the thesis evaluation strategy: regret and gap closed must be reported beside accuracy because the cost of a wrong label depends on the measured runtime distance from the oracle.")
    para(doc, "The high-regret errors have a different meaning. They indicate that the model crossed a boundary where the selected algorithm was not merely slightly slower, but materially slower. The expensive timsort_to_introsort cases suggest that some arrays still contain run-adaptive advantage, but the feature combination is not strong enough for the classifier to keep them inside the timsort region. The heapsort_to_introsort cases show the opposite problem: the model assigns a general-purpose introsort decision where the measured runtime landscape favours heapsort's more predictable behaviour.")
    figure(doc, figs.get("failure", Path()), "Figure X. Failed prediction pairs ordered by mean regret rather than by count alone.")

    heading(doc, "Worst-Failure Region", 2)
    wrows = []
    for _, r in worst.iterrows():
        wrows.append([r.pair, f"{int(r.n):,}", f"{r.mean_regret_us:.2f}", f"{r.median_n:.0f}", f"{r.median_margin_us:.2f}"])
    table(doc, ["Pair", "Cases", "Mean regret (us)", "Median n", "Median margin (us)"], wrows)
    srows = []
    for _, r in shift.iterrows():
        srows.append([r.feature, f"{r.worst_median:.4g}", f"{r.all_failure_median:.4g}", f"{r.all_data_median:.4g}"])
    table(doc, ["Feature", "Worst failures", "All failures", "All data"], srows)
    para(doc, "The worst 500 failures are structurally different from ordinary failures. Their median size is much larger than both the full dataset and the full failure set. They also have weaker longest-run structure, higher inversion ratio, and lower duplicate/frequency concentration. This combination creates a difficult region: the arrays are large enough for a wrong constant-factor choice to matter, but their structural evidence is not cleanly aligned with one algorithm.")
    para(doc, "This does not prove that a single feature causes failure. The correct interpretation is empirical concentration: the highest regret cases are concentrated in a region with larger arrays, less reliable run evidence, and weaker repetition signals. A future controlled experiment could isolate these variables synthetically, but the present evidence is sufficient to identify where the production model should be examined first.")

    heading(doc, "Domain Interpretation", 2)
    drows = []
    for _, r in domain.iterrows():
        drows.append([r.domain, f"{int(r.n):,}", r.dominant, f"{r.dominant_pct:.1f}%", f"{r.success_pct:.2f}%", f"{r.mean_regret_us:.2f}", f"{r.p95_regret_us:.2f}"])
    table(doc, ["Domain", "Rows", "Dominant", "Dominance", "Success", "Mean regret", "P95 regret"], drows)
    para(doc, "The domain results show that the same final portfolio does not behave identically across all sources. Crypto, Stock, and Earthquake remain timsort-dominated, while Weather is mixed and has substantially lower success. This does not contradict the selector's general result. Instead, it shows that domain composition changes the density of arrays near the decision boundary. Weather contributes fewer rows and a more mixed winner distribution, so its label-level accuracy is weaker even when many mistakes remain low-cost.")
    figure(doc, figs.get("domain", Path()), "Figure X. Domain-level winner distribution for the final three-algorithm problem.")

    heading(doc, "Interpretation for the Thesis", 2)
    para(doc, "The analysis strengthens the thesis argument in three ways. First, the final three algorithms are not interchangeable labels: timsort, heapsort, and introsort occupy different structural regions of the measured runtime landscape. Second, the classifier succeeds when the structural trace of the winning algorithm is visible in the sixteen-feature vector, especially for timsort's order/run-adaptive regime and heapsort's duplicate/frequency-concentrated regime. Third, failures must be interpreted through regret because many label errors occur where runtimes are nearly tied.")
    para(doc, "The remaining limitation is the boundary between the non-timsort algorithms and the high-regret tail of timsort errors. These regions likely require either richer structural features or lightweight proxies for low-level execution effects such as memory access and branch behaviour. The present results should therefore be claimed as an empirical behavioural analysis, not as a causal proof. They are strong enough to explain why the selector works and where it fails, while still leaving a clear path for future controlled experiments.")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
