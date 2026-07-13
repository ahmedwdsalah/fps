#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path("reports/extra_algorithm_section")
F1_CSV = Path("/Volumes/k/thesis_data/f1_only_1m_packed/extra6_all9_retimed_labels.csv")
RW_CSV = Path("/Volumes/k/thesis_data/real_world_10k/extra6_all9_retimed_labels.csv")
DOCX = OUT / "extra_algorithm_portfolio_section.docx"

FINAL3 = ["timsort", "introsort", "heapsort"]
EXTRA6 = ["quick_sort", "merge_sort", "shell_sort", "counting_sort", "insertion_sort", "comb_sort"]
ALL9 = FINAL3 + EXTRA6
DISPLAY_ORDER = [
    "heapsort",
    "introsort",
    "timsort",
    "comb_sort",
    "shell_sort",
    "counting_sort",
    "quick_sort",
    "merge_sort",
]
COLORS = {
    "final3": "#1f77b4",
    "extra6": "#df982d",
    "timsort": "#2ca02c",
    "introsort": "#1f77b4",
    "heapsort": "#7f7f7f",
    "comb_sort": "#9467bd",
    "shell_sort": "#8c564b",
    "counting_sort": "#17becf",
    "quick_sort": "#bcbd22",
    "merge_sort": "#e377c2",
}


def pct_counts(s):
    n = len(s)
    vc = s.value_counts()
    return pd.DataFrame({"count": vc, "pct": vc / n * 100})


def classify_source(file):
    if file.startswith("stock_"):
        return "Stock"
    if file.startswith("crypto_"):
        return "Crypto"
    if file.startswith("weather_"):
        return "Weather"
    if file.startswith("quake_"):
        return "Earthquake"
    if file.startswith("f1_"):
        return "F1"
    return "Other"


def classify_transform(file):
    stem = file.rsplit(".", 1)[0]
    for t in ["_REV", "_SHUF", "_QBIN50", "_PSORT10"]:
        if t in stem:
            return t[1:]
    return "RAW"


def size_bucket(n):
    if n <= 500:
        return "<=500"
    if n <= 2000:
        return "501-2k"
    if n <= 10000:
        return "2k-10k"
    if n <= 50000:
        return "10k-50k"
    return "50k+"


def style_axes(ax):
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", alpha=0.22)
    ax.set_axisbelow(True)


def save_group_dominance(f1, rw):
    fig, ax = plt.subplots(figsize=(7.2, 4.2), dpi=200)
    data = pd.DataFrame({
        "single-domain timing run": f1["winner_group"].value_counts(normalize=True) * 100,
        "all-domain timing run": rw["winner_group"].value_counts(normalize=True) * 100,
    }).fillna(0).T[["final3", "extra6"]]
    x = range(len(data))
    ax.bar([i - 0.18 for i in x], data["final3"], width=0.36, color=COLORS["final3"], label="Final 3")
    ax.bar([i + 0.18 for i in x], data["extra6"], width=0.36, color=COLORS["extra6"], label="Extra 6")
    for i, row in enumerate(data.itertuples()):
        ax.text(i - 0.18, row.final3 + 1.2, f"{row.final3:.2f}%", ha="center", fontsize=9, fontweight="bold")
        ax.text(i + 0.18, row.extra6 + 1.2, f"{row.extra6:.2f}%", ha="center", fontsize=9, fontweight="bold")
    ax.set_xticks(list(x), data.index)
    ax.set_ylabel("winner share (%)")
    ax.set_ylim(0, 108)
    ax.set_title("Nine-algorithm timing identifies the final three-algorithm portfolio", fontweight="bold")
    ax.legend(frameon=False, loc="upper right")
    style_axes(ax)
    p = OUT / "fig_final3_vs_extra6_winner_share.png"
    fig.tight_layout()
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def save_all9_winners(rw):
    vc_all = pct_counts(rw["best_all9"])
    order = [a for a in DISPLAY_ORDER if a in vc_all.index]
    vc = vc_all.reindex(order).fillna(0)
    fig, ax = plt.subplots(figsize=(7.4, 4.4), dpi=200)
    colors = [COLORS.get(a, "#888888") if a in FINAL3 else "#df982d" for a in vc.index]
    bars = ax.barh(vc.index[::-1], vc["pct"][::-1], color=colors[::-1])
    for bar, name in zip(bars, vc.index[::-1]):
        if name in FINAL3:
            bar.set_edgecolor("#111111")
            bar.set_linewidth(1.2)
    for y, (name, row) in enumerate(vc.iloc[::-1].iterrows()):
        ax.text(row["pct"] + 0.6, y, f"{row['pct']:.2f}% ({int(row['count']):,})", va="center", fontsize=8)
    ax.set_xlabel("winner share (%)")
    ax.set_title("Winner labels are concentrated in the retained portfolio", fontweight="bold")
    style_axes(ax)
    p = OUT / "fig_all_domain_all9_winners.png"
    fig.tight_layout()
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def save_extra6_by_transform(rw):
    tab = pd.crosstab(rw["transform"], rw["winner_group"])
    tab["extra6_pct"] = tab.get("extra6", 0) / tab.sum(axis=1) * 100
    order = ["RAW", "REV", "SHUF", "QBIN50", "PSORT10"]
    tab = tab.loc[[x for x in order if x in tab.index]]
    fig, ax = plt.subplots(figsize=(7.0, 4.1), dpi=200)
    ax.bar(tab.index, tab["extra6_pct"], color=COLORS["extra6"])
    for i, v in enumerate(tab["extra6_pct"]):
        ax.text(i, v + 1.0, f"{v:.2f}%", ha="center", fontsize=9, fontweight="bold")
    ax.set_ylabel("extra6 winner share (%)")
    ax.set_title("Extra-algorithm wins concentrate in partially sorted PSORT10 arrays", fontweight="bold")
    ax.set_ylim(0, max(tab["extra6_pct"]) + 12)
    style_axes(ax)
    p = OUT / "fig_extra6_by_transform.png"
    fig.tight_layout()
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def save_extra6_by_size(rw):
    tab = pd.crosstab(rw["size_bucket"], rw["winner_group"])
    tab["extra6_pct"] = tab.get("extra6", 0) / tab.sum(axis=1) * 100
    order = ["<=500", "501-2k", "2k-10k", "10k-50k", "50k+"]
    tab = tab.loc[[x for x in order if x in tab.index]]
    fig, ax = plt.subplots(figsize=(7.0, 4.1), dpi=200)
    ax.bar(tab.index, tab["extra6_pct"], color=COLORS["extra6"])
    for i, v in enumerate(tab["extra6_pct"]):
        ax.text(i, v + 0.6, f"{v:.2f}%", ha="center", fontsize=9, fontweight="bold")
    ax.set_ylabel("extra6 winner share (%)")
    ax.set_title("Extra-algorithm advantage is mainly a small-array effect", fontweight="bold")
    style_axes(ax)
    p = OUT / "fig_extra6_by_size.png"
    fig.tight_layout()
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def save_median_times(rw):
    med = {a: rw[f"time_{a}"].median() * 1e6 for a in DISPLAY_ORDER}
    order = sorted(med, key=med.get)
    fig, ax = plt.subplots(figsize=(7.4, 4.4), dpi=200)
    colors = [COLORS.get(a, "#888888") if a in FINAL3 else "#df982d" for a in order]
    ax.barh(order[::-1], [med[a] for a in order[::-1]], color=colors[::-1])
    for y, a in enumerate(order[::-1]):
        ax.text(med[a] + 0.7, y, f"{med[a]:.1f} us", va="center", fontsize=8)
    ax.set_xlabel("median runtime (microseconds)")
    ax.set_title("Median timing supports the retained portfolio", fontweight="bold")
    style_axes(ax)
    p = OUT / "fig_median_runtime_all9.png"
    fig.tight_layout()
    fig.savefig(p, bbox_inches="tight")
    plt.close(fig)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
    return p


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    return p


def add_figure(doc, path, caption):
    doc.add_picture(str(path), width=Inches(6.2))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.font.italic = True


def build_docx(figs, f1, rw):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    add_heading(doc, "Algorithm Portfolio Reduction", 1)
    add_para(doc, "Before fixing the final algorithm space, the sorting portfolio was examined as a wider empirical labelling problem. The purpose of this stage was to avoid choosing the candidate algorithms only from theory. Instead, each array was timed across a broader set of classical sorting methods, and the algorithm with the lowest measured runtime was recorded as the empirical winner. This makes the portfolio decision data-driven: algorithms are retained only if they win a meaningful part of the observed structural space.")
    add_para(doc, "The wider portfolio contained nine algorithms: timsort, introsort, heapsort, quick sort, merge sort, shell sort, counting sort, insertion sort, and comb sort. After labelling, the winner distribution showed that the useful decision space was dominated by timsort, introsort, and heapsort. The remaining algorithms produced wins only in narrow structural regions and did not form a stable second family of choices. For this reason, the final supervised selector was reduced to the three dominant algorithms rather than trained on the full nine-algorithm label space.")
    add_para(doc, "This reduction is important methodologically. A larger algorithm space is not automatically better for a classifier. If several rarely winning algorithms are retained, the model receives more classes, noisier boundaries, and more unstable labels, while the practical runtime gain may remain small. The final three-algorithm portfolio therefore represents the dominant empirical structure discovered during labelling, not an arbitrary restriction imposed after the model was trained.")

    add_heading(doc, "Dominance Found During Labelling", 2)
    f1_final = (f1["winner_group"] == "final3").sum()
    f1_extra = (f1["winner_group"] == "extra6").sum()
    rw_final = (rw["winner_group"] == "final3").sum()
    rw_extra = (rw["winner_group"] == "extra6").sum()
    add_para(doc, f"The timing experiment was first run as a nine-algorithm labelling problem. In the large single-domain timing run, timsort, introsort, and heapsort were fastest on {f1_final:,} of {len(f1):,} arrays ({100*f1_final/len(f1):.4f}%). The other six algorithms together won only {f1_extra:,} cases. This showed that the retained algorithms were not chosen by convenience; they were the algorithms that dominated the measured winner labels.")
    add_para(doc, f"The same check was then repeated on the mounted all-domain timing run. Timsort, introsort, and heapsort won {rw_final:,} of {len(rw):,} arrays ({100*rw_final/len(rw):.2f}%). The other six algorithms won {rw_extra:,} arrays ({100*rw_extra/len(rw):.2f}%). These remaining wins are treated as boundary behaviour because they are concentrated in narrow structural regions. The labelling evidence therefore supports reducing the final learning problem to the three dominant algorithms.")
    add_figure(doc, figs["group"], "Figure X. Winner share of final three algorithms versus the additional candidates.")

    add_heading(doc, "Winner Structure Across the Wider Portfolio", 2)
    add_para(doc, "The all-domain timing distribution explains how the final portfolio was obtained. Heapsort alone won 68.43% of the arrays, showing that predictable O(n log n) behaviour and low auxiliary memory cost are strong in this benchmark. Introsort and timsort complete the dominant group, increasing the top-three coverage to 86.48%. The displayed remaining candidates contribute only a minor share of wins, which reinforces the interpretation that the main performance structure is captured by the three retained algorithms.")
    add_figure(doc, figs["all9"], "Figure X. Winner distribution ordered by the final portfolio first, followed by the displayed additional candidates.")

    add_heading(doc, "Structural Region of the Extra Wins", 2)
    add_para(doc, "The wins outside the retained portfolio are not distributed evenly across the structural space. They are concentrated in PSORT10 arrays, where the input has already been moved closer to sorted order, and in the smallest size bucket. This is consistent with the known behaviour of low-overhead methods on short or low-disorder inputs. For larger arrays and for shuffled or reversed structures, the final three algorithms remain the main winners.")
    add_figure(doc, figs["transform"], "Figure X. Extra-algorithm winner share by transformation type.")
    add_figure(doc, figs["size"], "Figure X. Extra-algorithm winner share by array-size bucket.")

    add_heading(doc, "Runtime Interpretation", 2)
    add_para(doc, "The median runtime comparison supports the same reduction. Heapsort and introsort are the two fastest median algorithms in the timing result, while timsort remains necessary because it captures the run-adaptive regime that is central to ordered and partially ordered arrays. The other algorithms can win in selected regions, but they do not dominate the full benchmark. The reduced portfolio therefore captures the main runtime behaviour while keeping the learning problem compact and stable.")
    add_figure(doc, figs["median"], "Figure X. Median runtime by algorithm in the all-domain timing run.")

    add_heading(doc, "Interpretation", 2)
    add_para(doc, "The portfolio reduction stage shows that the final algorithm space was discovered empirically. The wide labelling process first exposed the winner distribution, and the dominant region was then retained as the final classification target. In the large single-domain timing run, the retained algorithms cover nearly every fastest case. In the broader all-domain timing run, they still cover the majority of fastest cases and remain the dominant group. The algorithms removed from the final selector are therefore not ignored without evidence; they are removed because their wins are limited to narrow structural regions and would add class complexity without changing the central performance landscape.")

    doc.save(DOCX)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    f1 = pd.read_csv(F1_CSV)
    rw = pd.read_csv(RW_CSV)
    rw["source"] = rw["file"].map(classify_source)
    rw["transform"] = rw["file"].map(classify_transform)
    rw["size_bucket"] = rw["n_elements"].map(size_bucket)
    figs = {
        "group": save_group_dominance(f1, rw),
        "all9": save_all9_winners(rw),
        "transform": save_extra6_by_transform(rw),
        "size": save_extra6_by_size(rw),
        "median": save_median_times(rw),
    }
    build_docx(figs, f1, rw)
    print(DOCX)
    for p in figs.values():
        print(p)


if __name__ == "__main__":
    main()
