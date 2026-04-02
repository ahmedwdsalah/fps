#!/usr/bin/env python3
"""
Generate Thesis Chapters 1-3 as a DOCX file following CIU Thesis Writing Guidelines.

Formatting rules (CIU):
- Font: Times New Roman, 12pt body, 14pt chapter titles
- Margins: 4cm top, 3.5cm left, 3cm right, 2.5cm bottom
- Line spacing: 1.5 throughout
- Headings: Arabic numerals, decimal system
- Chapter titles: 14pt, Bold, CAPITAL LETTERS, centered
- Section titles (Heading 2): 12pt, Bold, CAPITAL LETTERS
- Sub-section titles (Heading 3): 12pt, Bold, CAPITAL LETTERS
- No paragraph indentation, justified alignment
- Tables: Title above (bold number, not bold title), 10pt inside
- Figures: Title below, centered
- Equations: Centered, numbered right-aligned
- Citation: APA style
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)),
                           "thesis_full_draft.docx")


def set_margins(section):
    """Set CIU margins: 4cm top, 3.5cm left, 3cm right, 2.5cm bottom."""
    section.top_margin = Cm(4)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(3)
    section.bottom_margin = Cm(2.5)


def set_paragraph_format(paragraph, space_after=Pt(0), space_before=Pt(0),
                         line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Set standard CIU paragraph formatting."""
    pf = paragraph.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = line_spacing
    pf.alignment = alignment
    pf.first_line_indent = Cm(0)


def add_chapter_title(doc, chapter_num, title):
    """Add chapter heading: 14pt, Bold, CAPS, centered."""
    # "CHAPTER ONE" line (skip for unnumbered sections like References)
    chapter_words = {1: "ONE", 2: "TWO", 3: "THREE", 4: "FOUR", 5: "FIVE"}
    if chapter_num is not None:
        p = doc.add_paragraph()
        set_paragraph_format(p, space_before=Pt(24), space_after=Pt(12),
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(f"CHAPTER {chapter_words[chapter_num]}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Chapter title line
    p2 = doc.add_paragraph()
    set_paragraph_format(p2, space_before=Pt(6), space_after=Pt(18),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run2 = p2.add_run(title.upper())
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.name = "Times New Roman"
    run2.font.color.rgb = RGBColor(0, 0, 0)


def add_section_heading(doc, number, title):
    """Add section heading (Heading 2): 12pt, Bold, CAPS."""
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(18), space_after=Pt(6),
                         alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(f"{number} {title.upper()}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_subsection_heading(doc, number, title):
    """Add sub-section heading (Heading 3): 12pt, Bold, CAPS."""
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12), space_after=Pt(6),
                         alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(f"{number} {title.upper()}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_body_text(doc, text, space_after=Pt(6)):
    """Add body paragraph: 12pt, Times New Roman, 1.5 spacing, justified, no indent."""
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=space_after)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_italic_text(doc, text):
    """Add italic body paragraph."""
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(6))
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.italic = True
    return p


def add_equation(doc, equation_text, eq_number):
    """Add centered equation with right-aligned number."""
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(6), space_after=Pt(6),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(equation_text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.italic = True
    # Add tab and equation number
    run2 = p.add_run(f"    ({eq_number})")
    run2.font.size = Pt(12)
    run2.font.name = "Times New Roman"
    return p


def add_table(doc, title, headers, rows, table_number):
    """Add a table with CIU formatting."""
    # Table title
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12), space_after=Pt(6),
                         alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run_bold = p.add_run(f"Table {table_number}: ")
    run_bold.bold = True
    run_bold.font.size = Pt(12)
    run_bold.font.name = "Times New Roman"
    run_title = p.add_run(title)
    run_title.font.size = Pt(12)
    run_title.font.name = "Times New Roman"

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add borders (top header and bottom)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    return table


def add_citation(text, author, year):
    """Format an in-text APA citation."""
    return f"{text} ({author}, {year})"


# ===========================================================================
# CHAPTER CONTENT
# ===========================================================================

def write_chapter_1(doc):
    """Chapter 1: Introduction."""
    add_chapter_title(doc, 1, "Introduction")

    # 1.1 Background
    add_section_heading(doc, "1.1", "Background")

    add_body_text(doc,
        "Sorting sits right at the heart of computer science — it is, arguably, the most studied "
        "algorithmic problem in the field. Databases rely on it, machine learning pipelines lean on it, "
        "and on any given day, sorting routines fire off billions of times across the world's servers. "
        "Yet here is the thing that still catches people off guard: after all these decades, there is no "
        "single sorting algorithm that wins everywhere. How well a particular sort performs hinges on "
        "properties of the data itself — the array's length, how much existing order it already has, whether "
        "values repeat heavily, the statistical shape of the distribution (Rice, 1976). It is deeply "
        "context-dependent."
    )

    add_body_text(doc,
        "Take timsort. It is the default sort in Python and Java, and the reason is fascinating: timsort "
        "hunts for natural runs — stretches of data that are already in order — and merges them cleverly, "
        "so partially-sorted inputs fly through in nearly linear time. Introsort, on the other hand (the "
        "default in C++ standard libraries and behind NumPy's np.sort), starts with quicksort but bails "
        "out to heapsort if the recursion gets too deep, which guarantees O(n log n) worst-case. Then "
        "there is heapsort itself: rock-solid O(n log n) always, O(1) extra memory, but its tree-indexed "
        "access pattern hammers the CPU cache. Each one shines under different conditions and stumbles "
        "under others."
    )

    add_body_text(doc,
        "A real example drives this home. Imagine a nearly-sorted financial time series — say 100,000 "
        "price ticks that are almost in ascending order already. Timsort chews through it roughly 10 "
        "times faster than heapsort, because it spots those long ascending runs and merges them in "
        "essentially one pass. Now shuffle that same array randomly, and suddenly introsort is king, "
        "thanks to quicksort's cache-friendly scanning pattern on unstructured data. A data pipeline "
        "that blindly picks one sorting algorithm every time? It is leaving real performance on the table, "
        "and quite a lot of it."
    )

    add_body_text(doc,
        "Rice (1976) gave this dilemma a formal name: the Algorithm Selection Problem. The question is "
        "deceptively simple — given a problem instance, which algorithm from your toolbox should you pick "
        "to get the best performance? Researchers have tackled this successfully for SAT solving "
        "(Xu et al., 2008), combinatorial optimisation (Kotthoff, 2016), even ML model selection "
        "(Brazdil et al., 2009). But for sorting? Surprisingly, runtime sorting selection with modern "
        "ML approaches and rigorous evaluation remains largely uncharted territory."
    )

    add_body_text(doc,
        "There is a further wrinkle, and it matters a lot in practice. Production pipelines do not sit "
        "still. A system trained on financial tick data might suddenly start ingesting IoT sensor "
        "readings or server log files — data with completely different structural fingerprints. A "
        "classifier frozen at training time simply cannot cope with that kind of drift. Which is "
        "exactly why this work incorporates an online learning component: something that keeps "
        "refining its decisions as new, potentially unfamiliar data rolls in, no full retraining needed."
    )

    # 1.2 Problem Statement
    add_section_heading(doc, "1.2", "Problem Statement")

    add_body_text(doc,
        "The problem of algorithm selection is fundamentally about mismatch: when a computational task can be solved by multiple algorithms "
        "with vastly different performance profiles across varying input domains, a single default choice leaves performance on the table. Rather "
        "than accept that fixed default—adequate on average but suboptimal nearly everywhere—a selector could predict which algorithm to run for "
        "each specific input, trading a small overhead for larger gains. Yet this trade-off only works if three brutal constraints are satisfied "
        "simultaneously, and these constraints are in genuine tension with one another. The first is computational efficiency: the cost of computing "
        "features and making predictions must be lower than the speedup achieved; there is no room for offline retraining or expensive feature extraction—"
        "the selector must be lightweight and work at deployment time. The second is generalisation without retraining: production systems encounter data "
        "distributions that drift away from training data; you have no way to know what patterns will arrive, yet the selector must adapt or at least not "
        "collapse catastrophically. The third is latency—predictions must be instantaneous. In practice, these three constraints pull against each other: "
        "more sophisticated models have higher computational cost, they tend to overfit to specific domains, and they have longer inference times. The problem, "
        "then, is not just to predict which algorithm is fastest, but to do so while respecting all three constraints."
    )

    add_body_text(doc,
        "Formally, the selection problem can be stated as follows. Given a set of inputs X and a portfolio of algorithms A = {a_1, a_2, ..., a_k}, define "
        "T_a(x) as the execution time of algorithm a on input x. The oracle selector a*(x) = argmin_a T_a(x) always picks the fastest algorithm; but "
        "computing the oracle requires observing all execution times, which defeats the purpose. A practical selector s(f(x)) instead computes features f(x) "
        "from the input and uses those to predict which algorithm to run. The objective is to learn s such that the expected regret is minimised:"
    )

    add_equation(doc,
        "Regret(s) = E[ T_{s(f(x))}(x) - T_{a*(x)}(x) ]",
        "1.1"
    )

    add_body_text(doc,
        "where the expectation is taken over the deployment distribution of inputs. The regret quantifies how much worse the selector's choice is compared to "
        "the oracle on average. For selection to deliver value, the full end-to-end cost—feature computation plus execution of the selected algorithm—must be "
        "lower than simply running a fixed default algorithm every time. Accuracy alone is meaningless if the overhead eats the savings."
    )

    add_body_text(doc,
        "A final but critical point concerns evaluation. Algorithm selection research suffers from a widespread and often unintentional bias: when training and "
        "test data come from the same source or benchmark, models learn source-specific quirks—preprocessing artefacts, data collection idiosyncrasies, systematic "
        "biases—rather than genuinely transferable structural patterns. Reported accuracy looks impressive, but generalisation to new sources collapses. To avoid "
        "this trap requires strict source-level separation at the train-test boundary, ensuring that any measured accuracy reflects genuine out-of-distribution "
        "performance, not memorisation of source-level noise."
    )

    # 1.3 Motivation
    add_section_heading(doc, "1.3", "Motivation")

    add_body_text(doc,
        "Three things came together to motivate this work. The first, and perhaps most compelling: the "
        "VBS-SBS gap is not trivial. Measurements on diverse real-world datasets show that sticking "
        "with one fixed algorithm wastes a meaningful chunk of total sorting time. On arrays with "
        "pronounced structural patterns — like nearly-sorted financial series or sensor data full of "
        "repeated values — choosing the wrong sort can cost you dearly, sometimes by an order of magnitude."
    )

    add_body_text(doc,
        "The second motivation: extracting structural features is cheap. Really cheap. Things like the "
        "fraction of already-sorted neighbours, how many monotonic runs exist, duplicate concentration, "
        "entropy — you can grab all of these in one O(n) sweep. And they carry strong predictive signal, "
        "particularly when it comes to telling apart timsort-friendly inputs (lots of pre-existing order) "
        "from inputs where a straightforward comparison sort does better (random or high-duplication data)."
    )

    add_body_text(doc,
        "Third, in deployment environments such as data engineering pipelines, the input distribution may "
        "shift over time. A static classifier trained once? It cannot keep up. That is precisely what "
        "drives the two-layer design in this thesis: an offline classifier gives you a solid starting "
        "policy right away, while an online contextual bandit keeps adjusting the picks as new, "
        "unfamiliar data streams in. Exploration versus exploitation, handled gracefully. The cold-start "
        "problem and the non-stationarity of real workloads — both addressed in one architecture."
    )

    # 1.4 Research Questions
    add_section_heading(doc, "1.4", "Research Questions")

    add_body_text(doc,
        "Four research questions guide this thesis:"
    )

    add_body_text(doc,
        "RQ1: Is a compact set of O(n) structural features enough to meaningfully capture how "
        "different sorting algorithms perform on real-world data?"
    )

    add_body_text(doc,
        "RQ2: What accuracy ceiling can a machine learning classifier realistically hit when choosing "
        "among three C-level sorting implementations (introsort, heapsort, timsort) on diverse inputs?"
    )

    add_body_text(doc,
        "RQ3: Is the sorting algorithm selection problem fundamentally binary — timsort versus "
        "comparison-sort — rather than a genuine three-way classification?"
    )

    add_body_text(doc,
        "RQ4: Can a contextual bandit adapt its selection policy on the fly when the data distribution "
        "drifts away from what it was trained on?"
    )

    # 1.5 Objectives
    add_section_heading(doc, "1.5", "Objectives")

    add_body_text(doc,
        "This thesis pursues six concrete objectives:"
    )

    add_body_text(doc,
        "1. Design and validate 16 structural features — all computable in O(n) — that capture the "
        "array properties most relevant to sorting algorithm performance."
    )

    add_body_text(doc,
        "2. Assemble a genuinely diverse real-world dataset spanning finance, ecology, network traffic, "
        "medical records, geospatial coordinates, text, and public ML benchmarks, with strict source-level "
        "separation so that no data leakage can creep in."
    )

    add_body_text(doc,
        "3. Train an XGBoost gradient-boosted tree classifier for offline sorting algorithm selection "
        "and report honest metrics on held-out test data with zero source leakage between splits."
    )

    add_body_text(doc,
        "4. Diagnose the theoretical ceiling of three-class classification by measuring just how "
        "similar introsort and heapsort really are on real-world data."
    )

    add_body_text(doc,
        "5. Design and implement a LinUCB contextual bandit for online adaptation to distribution shift."
    )

    add_body_text(doc,
        "6. Package the system as a deployable Python library with a simple API."
    )

    # 1.6 Thesis Contributions
    add_section_heading(doc, "1.6", "Thesis Contributions")

    add_body_text(doc,
        "The contributions of this work are as follows:"
    )

    add_body_text(doc,
        "1. Feature engineering for sorting. Sixteen O(n) features organised into groups covering scale, "
        "sortedness, uniqueness, spread, structure, randomness, shape, outliers, and local order. Every "
        "feature is formally defined and validated against real-world data — not just synthetic benchmarks."
    )

    add_body_text(doc,
        "2. A large-scale empirical performance map. Three C-level sorting implementations (NumPy) "
        "benchmarked across 37,976 real-world arrays from 9 domains. The evaluation methodology — "
        "GroupShuffleSplit by source, a 5% tie filter, zero bootstrapping — is designed to be as "
        "honest as possible."
    )

    add_body_text(doc,
        "3. A ceiling analysis that actually explains the results. The diagnostic reveals that "
        "3-class accuracy tops out around 71% because introsort and heapsort have a median timing "
        "gap of just 0.8%. Collapse those two into one class and binary accuracy jumps to 94.2%. "
        "That finding, on its own, reshapes how we should think about sorting algorithm selection."
    )

    add_body_text(doc,
        "4. A two-layer adaptive selection system. The offline XGBoost classifier handles cold-start, "
        "and the online LinUCB contextual bandit picks up from there, adapting to distribution shift "
        "without a single retraining cycle."
    )

    add_body_text(doc,
        "5. Full reproducibility. Every script, every data collection pipeline, every trained model "
        "artifact and evaluation result is published with provenance tracking, so anyone can verify "
        "or build on this work."
    )

    # 1.7 Thesis Outline
    add_section_heading(doc, "1.7", "Thesis Outline")

    add_body_text(doc,
        "The rest of this document is laid out as follows. Chapter 2 covers the theoretical background "
        "and related work — sorting algorithms, the ASP framework, per-instance selection systems, "
        "gradient-boosted trees, and contextual bandits. Chapter 3 dives into the methodology: system "
        "architecture, feature engineering, data collection, model training, and how the evaluation was "
        "set up. Chapter 4 is where the numbers live — classification results, regret analysis, ceiling "
        "diagnosis, and baseline comparisons. Chapter 5 wraps things up with a discussion of what it "
        "all means, the limitations, and where future work could take this."
    )


def write_chapter_2(doc):
    """Chapter 2: Background and Related Work."""
    doc.add_page_break()
    add_chapter_title(doc, 2, "Background and Related Work")

    # 2.1 Sorting Algorithms
    add_section_heading(doc, "2.1", "Sorting Algorithms")

    add_body_text(doc,
        "At the most basic level, a sorting algorithm rearranges elements into some specified order. "
        "For comparison-based sorts — the family this thesis works with — all ordering information comes "
        "from pairwise comparisons, and there is a well-known lower bound: any such algorithm needs "
        "at least O(n log n) comparisons in the worst case (Knuth, 1998). The reasoning is elegant, "
        "actually. With n! possible permutations to distinguish between, you need log2(n!) binary "
        "decisions at minimum, and that works out to O(n log n). All three algorithms considered here "
        "hit that bound asymptotically, but their constant factors, how they treat the CPU cache, and "
        "whether they can exploit pre-existing structure in the data — those details vary enormously."
    )

    add_body_text(doc,
        "This thesis sticks to three comparison-based algorithms that ship as C-level implementations "
        "inside NumPy: introsort, heapsort, and timsort. That restriction is intentional and important. "
        "By keeping everything at the same implementation level, any performance difference the model "
        "picks up on reflects genuine algorithmic behaviour, not language overhead or coding quality. "
        "If you mixed, say, a hand-written Python mergesort with a compiled C introsort, the model "
        "would just learn that C is faster than Python — not terribly useful."
    )

    # 2.1.1 Introsort
    add_subsection_heading(doc, "2.1.1", "Introsort")

    add_body_text(doc,
        "Musser (1997) came up with introsort as a way to have your cake and eat it too. It kicks off "
        "with quicksort — which is fast in practice thanks to cache-friendly sequential scans — but "
        "monitors the recursion depth. Once that depth exceeds 2 * floor(log2(n)), it bails out to "
        "heapsort, guaranteeing O(n log n) worst-case. In NumPy, you get introsort by calling "
        "numpy.sort(kind='quicksort'), and it is the default sorting algorithm."
    )

    add_body_text(doc,
        "During the quicksort phase, pivot selection uses a median-of-three heuristic: pick the median "
        "of the first, middle, and last elements. That simple trick avoids the dreaded O(n^2) worst case "
        "you get from naive pivot selection on sorted inputs, while still giving decent partition balance "
        "on average. The partitioning itself scans inward from both ends of the array, which keeps the "
        "working set in the CPU cache nicely. For tiny subarrays (typically under 16 elements), it drops "
        "down to insertion sort, whose low constant factors make it faster at small n."
    )

    add_body_text(doc,
        "Space-wise, introsort uses O(log n) for its recursion stack. On average, it does about "
        "1.39 * n * log2(n) comparisons, which is highly competitive on random data. The catch — "
        "and it matters for our story — is that introsort is completely blind to pre-existing order. "
        "Hand it a nearly-sorted array and it will process it the same way it processes a totally "
        "random permutation. No shortcuts."
    )

    # 2.1.2 Heapsort
    add_subsection_heading(doc, "2.1.2", "Heapsort")

    add_body_text(doc,
        "Heapsort (Williams, 1964) works in two distinct phases. First, it builds a max-heap in-place "
        "using Floyd's bottom-up construction — that takes O(n), which is surprisingly fast. Then it "
        "repeatedly pulls out the maximum element and re-heapifies, doing n - 1 extract-max operations "
        "at O(log n) each. Total: O(n log n) always. No best case, no worst case — just O(n log n) "
        "regardless of what the input looks like."
    )

    add_body_text(doc,
        "The big selling point of heapsort is memory: O(1) auxiliary space. It sorts completely "
        "in-place, no extra buffers needed. That makes it attractive for memory-tight environments. "
        "But — and this is the reason heapsort often feels sluggish in practice — its access pattern "
        "is terrible for modern caches. The heap's parent-child relationships jump around in memory "
        "rather than scanning sequentially, causing 2-3 times the cache misses you would see in "
        "quicksort (LaMarca and Ladner, 1999). On today's hardware, where cache performance "
        "often matters more than raw comparison counts, that penalty adds up."
    )

    add_body_text(doc,
        "Interestingly, despite the cache penalty, heapsort holds its own on data loaded with duplicate "
        "values. Quicksort-based algorithms can get tripped up by poor pivot selection when many elements "
        "are identical, leading to lopsided partitions. Heapsort just does not care. In NumPy, you call "
        "it with numpy.sort(kind='heapsort')."
    )

    # 2.1.3 Timsort
    add_subsection_heading(doc, "2.1.3", "Timsort")

    add_body_text(doc,
        "Tim Peters designed timsort in 2002 for Python's built-in sort, and it later found its way "
        "into Java's Arrays.sort for objects. The central idea is beautifully pragmatic: real-world "
        "data often contains natural runs — stretches that are already sorted or reverse-sorted — so "
        "why not exploit them? Timsort works in three phases. First, it scans through the input "
        "identifying ascending and descending runs, reversing descending ones in-place. Second, any run "
        "shorter than a minimum length (minrun, usually 32–64) gets padded out using binary insertion "
        "sort. Third, the accumulated runs are merged according to a carefully crafted policy that "
        "maintains a stack invariant and guarantees O(n log n) merge operations in total."
    )

    add_body_text(doc,
        "What really makes timsort clever is its galloping merge. When merging two runs and one "
        "side keeps winning the comparisons, timsort switches from plodding linear merge to "
        "exponential search — galloping, as Peters named it — which can leap over large blocks in "
        "O(log k) time. On nearly-sorted data with long runs, the payoff is massive."
    )

    add_body_text(doc,
        "Feed timsort an already-sorted array and it is done in O(n) — one scan, that is it. With "
        "k natural runs, the complexity is O(n log k), which is optimal for adaptive sorting (Barbay "
        "and Navarro, 2013). But throw completely random data at it and the overhead of run detection "
        "plus the galloping bookkeeping costs roughly 30–50% more comparisons than introsort. Timsort "
        "also needs O(n) extra memory for its merge buffer, which is a non-trivial additional cost. "
        "In NumPy, timsort goes by numpy.sort(kind='stable'), and yes, it does guarantee stability — "
        "equal elements preserve their original relative order."
    )

    add_body_text(doc,
        "So there is a natural landscape here for picking algorithms. Timsort shines when data has "
        "structure — sorted chunks, near-sortedness, long runs. Introsort and heapsort duke it out on "
        "random or heavily duplicated data. And here is the spoiler that Chapter 4 will unpack in "
        "detail: on the majority of real-world arrays, introsort and heapsort are practically "
        "indistinguishable. The median timing gap between them? A mere 0.8%."
    )

    # 2.1.4 Dropped Algorithms
    add_subsection_heading(doc, "2.1.4", "Non-Comparison Sorts")

    add_body_text(doc,
        "Non-comparison sorts like counting sort and radix sort did get a look during the design stage, "
        "but they did not make the cut. Counting sort scored zero wins across 720 benchmark arrays — "
        "the value range of real-world floats (up to 30 million) makes the bincount allocation too "
        "expensive. Radix sort on floats, via IEEE 754 bit-manipulation, was 20 times slower than "
        "introsort in pilot runs because of the bit-level conversion overhead. Both would have needed "
        "Cython or C extensions to be competitive, which would break the thesis constraint that "
        "everything uses NumPy's native C-level sorts. Apples to apples, not apples to oranges."
    )

    # 2.2 The Algorithm Selection Problem
    add_section_heading(doc, "2.2", "The Algorithm Selection Problem")

    add_body_text(doc,
        "Rice (1976) set down the formal scaffolding for what we now call the Algorithm Selection "
        "Problem (ASP). The idea: given a problem space P, a feature space F, an algorithm space A, "
        "and a performance measure y, find a mapping S: F -> A that optimises expected performance. "
        "It sounds abstract, but in the five decades since, this formulation has been instantiated "
        "in domain after domain and remains the go-to framework for per-instance algorithm selection."
    )

    add_body_text(doc,
        "Breaking it down: P is every possible problem instance — in our case, every possible input "
        "array. F is a d-dimensional representation that captures what matters about each instance for "
        "predicting algorithm behaviour. A is the portfolio of available algorithms. And y maps each "
        "(instance, algorithm) pairing to a performance score (runtime, solution quality, whatever "
        "you are optimising for). The selection mapping S: F -> A is what you actually learn."
    )

    add_body_text(doc,
        "Translated to sorting: P is all possible numerical arrays, F = R^16 is the structural feature "
        "vector this thesis defines, A = {introsort, heapsort, timsort}, and y is wall-clock execution "
        "time. The mapping S gets learned from labelled instances where every array has been timed on "
        "all three algorithms and tagged with whichever was fastest. The VBS is the theoretically perfect "
        "S, while the SBS is the boring option of just always picking the single globally best algorithm."
    )

    add_body_text(doc,
        "The VBS-SBS gap\u2014defined as (T_SBS \u2212 T_VBS) / T_SBS\u2014puts a number on how much "
        "room there is for per-instance selection to help. If the gap is large, then picking algorithms "
        "per-input genuinely matters. If it is tiny, one algorithm handles everything fine and the whole "
        "enterprise is pointless. In SAT solving, gaps exceeding 50 per cent are common (Xu et al., 2008), "
        "which explains why that community built so many sophisticated selectors. In sorting, as this "
        "thesis reveals, the gap depends heavily on what the dataset looks like and which algorithms "
        "compose the portfolio."
    )

    # 2.3 Per-Instance Algorithm Selection
    add_section_heading(doc, "2.3", "Per-Instance Algorithm Selection Systems")

    add_body_text(doc,
        "Per-instance algorithm selection takes the ASP a step further: instead of picking one algorithm "
        "for everything, it makes a fresh decision for every single input. The system that really got "
        "this idea off the ground was SATzilla (Xu et al., 2008). It selects among SAT solvers by "
        "looking at 48 features pulled from each SAT instance\u2014things like the clause-to-variable "
        "ratio, graph properties from the clause-variable incidence structure, and even probing features "
        "from short preliminary solver runs. A cost-sensitive regression model learns to predict how "
        "long each solver will take, then picks the fastest one."
    )

    add_body_text(doc,
        "SATzilla essentially wrote the playbook everyone else follows: define an instance feature space, "
        "run all candidate algorithms on a diverse training set to gather performance data, train a "
        "prediction model (regression or classification, depending on taste), and then evaluate against "
        "VBS and SBS baselines. That is a four-step recipe. This thesis adopts the same recipe\u2014adapted, "
        "of course, to the sorting domain."
    )

    add_body_text(doc,
        "AutoFolio (Lindauer et al., 2015) generalised SATzilla into something more flexible\u2014a "
        "configurable framework that handles feature selection, model choice, and hyperparameter tuning "
        "through Bayesian optimisation, all automatically. The insight was that algorithm selection itself "
        "is a meta-learning problem; you co-optimise the portfolio and the selector at the same time. "
        "AutoFolio has been thrown at constraint satisfaction, graph colouring, mixed-integer programming, "
        "and in each case it beats the single best algorithm baseline."
    )

    add_body_text(doc,
        "Other frameworks have piled on since. LLAMA (Kotthoff, 2013) and AS (Bischl et al., 2016) "
        "bring k-nearest neighbours, random forests, and support vector machines into the model family "
        "mix. Meanwhile, the ASLib benchmark library (Bischl et al., 2016) standardised evaluation by "
        "packaging common instance feature sets and performance data for multiple problem domains into "
        "one consistent library. That made comparing systems a lot easier than it used to be."
    )

    add_body_text(doc,
        "In sorting specifically, Li and Mao (2009) tried decision trees to pick between algorithms using "
        "array properties\u2014size, presortedness, duplicate ratio. Simple structural features could predict "
        "performance, they showed, but their study only covered synthetic arrays with uniform and normal "
        "distributions. That is a narrow slice of reality. Learned Sort (Kristo et al., 2020) went in a "
        "different direction entirely, using a learned CDF model to predict where elements belong and "
        "achieve near-linear sorting time. This thesis occupies a complementary position: it does not "
        "invent a new sorting algorithm. Instead, it picks the best existing C-level implementation for "
        "whatever input happens to arrive."
    )

    # 2.4 Feature Engineering for Algorithm Selection
    add_section_heading(doc, "2.4", "Feature Engineering for Algorithm Selection")

    add_body_text(doc,
        "Good features make or break per-instance algorithm selection. There are three properties that "
        "matter. First, informative: the features have to actually correlate with performance differences "
        "between algorithms. Second, cheap: if computing the features takes longer than the time you save "
        "by making a better choice, the whole thing is self-defeating. Third, robust: features need to "
        "generalise across data domains, not just work on the training distribution."
    )

    add_body_text(doc,
        "Different domains have landed on different feature sets. In SAT solving, people compute the "
        "clause-to-variable ratio, graph properties of the clause structure, and probing features from "
        "short solver runs (Xu et al., 2008). Algorithm configuration work tends to use problem size, "
        "constraint density, and domain-specific structural properties (Hutter et al., 2011). The shared "
        "principle is always the same: capture enough structure that the model can tell instances apart."
    )

    add_body_text(doc,
        "For sorting, the relevant structural properties are fairly well understood thanks to decades of "
        "algorithm theory. Timsort thrives when it can exploit pre-existing order\u2014adjacent sorted ratio, "
        "inversion ratio, runs ratio, longest run ratio; call those the Sortedness group. Quicksort-based "
        "algorithms, on the other hand, benefit from good pivot selection, which comes more easily when "
        "values are spread out and duplicates are rare (Spread group: dispersion ratio, IQR normalised, "
        "MAD normalised). Then there is sheer array size, which determines the overhead balance between "
        "algorithms (Scale group: length normalised). These theoretical insights shaped the 16-feature set "
        "used in this thesis\u2014Chapter 3 spells out the details."
    )

    # 2.5 Gradient-Boosted Trees
    add_section_heading(doc, "2.5", "Gradient-Boosted Trees")

    add_body_text(doc,
        "Gradient-boosted decision trees\u2014GBDT for short\u2014work by stacking weak learners on top "
        "of one another. Each new shallow tree tries to fix whatever the ensemble so far still gets wrong, "
        "fitting the residual errors from the previous round. XGBoost (Chen and Guestrin, 2016) is probably "
        "the most widely used implementation; it adds regularisation terms to the objective function and "
        "employs histogram-based splits for faster training."
    )

    add_body_text(doc,
        "The objective XGBoost minimises at iteration t looks like this:"
    )

    add_equation(doc,
        "L(t) = sum_i l(y_i, y_hat_i^(t-1) + f_t(x_i)) + Omega(f_t)",
        "2.1"
    )

    add_body_text(doc,
        "where l is the loss function, f_t is the new tree, and Omega(f_t) = gamma * T + 0.5 * lambda * ||w||^2 "
        "penalises tree complexity. T counts the leaves, w holds the leaf weights, and gamma plus lambda are "
        "hyperparameters you tune. This regularisation matters a lot for algorithm selection specifically, "
        "because labels can be noisy\u2014when two algorithms finish within microseconds of each other, the "
        "winner is essentially decided by measurement jitter, and the model really should not overfit to that."
    )

    add_body_text(doc,
        "For multi-class problems, XGBoost offers the softmax objective (multi:softprob), which produces "
        "calibrated probability estimates for every class\u2014not just a hard prediction. The class with "
        "the highest probability wins. Early stopping on a held-out validation set keeps things honest: "
        "training halts the moment validation loss plateaus, preventing the model from memorising noise. "
        "The choice of multi:softprob over multi:softmax is deliberate here. Those calibrated probabilities "
        "get handed downstream to the contextual bandit as a warm-start prior, so they need to be real "
        "probabilities, not just argmax labels."
    )

    add_body_text(doc,
        "Under the hood, histogram-based tree construction (tree_method='hist') bins continuous features "
        "into at most 256 discrete buckets, then evaluates split candidates at bin boundaries rather than "
        "at every unique feature value. Complexity drops from O(n * d) to O(n_bins * d), where n_bins "
        "is far smaller than n. For the 16-feature sorting dataset with 20,633 training arrays, this "
        "shaves an order of magnitude off training time with barely any accuracy cost."
    )

    add_body_text(doc,
        "Why GBDT for algorithm selection? Several reasons. The models handle feature interactions "
        "naturally\u2014the interaction between array size and sortedness, for example, matters a great deal "
        "and tree ensembles capture that without manual feature engineering. They are robust to irrelevant "
        "features, they spit out feature importance rankings, and they train fast on tabular data. "
        "Kotthoff (2016) identified these same qualities as the reason XGBoost became the default in "
        "algorithm selection research. Compared to deep learning alternatives like neural networks with "
        "embedding layers, GBDT models need less data, train faster, and deliver comparable or better "
        "accuracy on tabular datasets with fewer than 100 features (Grinsztajn et al., 2022). With "
        "only 16 features in the sorting dataset, XGBoost was the obvious pick."
    )

    # 2.6 Contextual Bandits
    add_section_heading(doc, "2.6", "Contextual Bandits")

    add_body_text(doc,
        "Contextual bandits (Li et al., 2010) extend the classic multi-armed bandit by conditioning every "
        "decision on a context vector. Each round t looks like this: the agent sees a context x_t (the "
        "feature vector for the current array), picks an action a_t (a sorting algorithm), observes a "
        "reward r_t (negative runtime, or a binary flag for whether it chose the best algorithm), and "
        "then updates its policy. It is essentially learning what works while simultaneously trying to "
        "do well\u2014the classic exploration-exploitation tension."
    )

    add_body_text(doc,
        "LinUCB (Li et al., 2010) is one particular contextual bandit algorithm based on a simple "
        "assumption: the expected reward is a linear function of the context-action features. For each "
        "action a, LinUCB maintains a ridge regression model:"
    )

    add_equation(doc,
        "E[r_t | x_t, a_t = a] = x_t^T * theta_a",
        "2.2"
    )

    add_body_text(doc,
        "where theta_a is the parameter vector for action a. The clever part is how the algorithm picks "
        "actions\u2014it maximises an upper confidence bound:"
    )

    add_equation(doc,
        "a_t = argmax_a (x_t^T * theta_hat_a + alpha * sqrt(x_t^T * A_a^{-1} * x_t))",
        "2.3"
    )

    add_body_text(doc,
        "where A_a = I + sum_{s: a_s=a} x_s * x_s^T is the design matrix, theta_hat_a = A_a^{-1} * b_a "
        "is the regularised least-squares estimate, and alpha controls exploration-exploitation balance. "
        "That square root term is the exploration bonus\u2014it shrinks as more data accumulates for each "
        "action, which nudges the agent toward trying algorithms it knows less about. Over time, the bonus "
        "fades and the agent converges to pure exploitation."
    )

    add_body_text(doc,
        "Why contextual bandits rather than full reinforcement learning? Because the problem here is "
        "stateless. Choosing heapsort for one array does not change the distribution of future arrays. "
        "No temporal credit assignment, no state transitions, no Bellman equations\u2014just independent "
        "context-action-reward triples. That simplification makes the regret analysis tractable and keeps "
        "the implementation lightweight. LinUCB in particular offers closed-form updates (no gradient "
        "descent loops), theoretical regret guarantees, and the ability to warm-start from the offline "
        "XGBoost predictions. Pretty hard to beat for this use case."
    )

    add_body_text(doc,
        "The theoretical regret bound of LinUCB runs O(d * sqrt(T * ln(T))), where d is the feature "
        "dimensionality and T the number of rounds (Abbasi-Yadkori et al., 2011). With d = 16 features "
        "and practical horizons of 1,000 to 10,000 rounds, that bound implies pretty rapid convergence. "
        "What makes it especially appealing is that the bound holds even if the reward distribution shifts "
        "over time, provided the linear feature-reward relationship remains roughly valid. Array features "
        "and algorithm performance are governed by stable algorithmic properties\u2014O(n log n) does not "
        "suddenly become O(n^2) because it is Tuesday\u2014so that assumption holds up well."
    )

    add_body_text(doc,
        "An alternative worth mentioning is Thompson Sampling (Agrawal and Goyal, 2013), which maintains "
        "a full posterior distribution over the reward parameters and samples from it when choosing actions. "
        "Empirically, Thompson Sampling sometimes edges out LinUCB. But\u2014and this matters for sorting\u2014the "
        "Bayesian posterior update is heavier computationally, especially with multi-dimensional contexts. "
        "LinUCB's deterministic selection and closed-form updates fit the low-latency demands of a sorting "
        "selector much better. The selection overhead has to be far smaller than the sorting time itself, "
        "and LinUCB delivers on that."
    )

    # 2.7 Related Work
    add_section_heading(doc, "2.7", "Related Work in Algorithm Selection")

    add_body_text(doc,
        "Algorithm selection has been around in one form or another for decades, spanning SAT solving, "
        "scheduling, graph problems, and more. This section surveys the most relevant prior work, starting "
        "with the foundational systems and then drilling into three recent contributions that are closest "
        "to what this thesis does."
    )

    # 2.7.1 Foundational Algorithm Selection Systems
    add_subsection_heading(doc, "2.7.1", "Foundational Algorithm Selection Systems")

    add_body_text(doc,
        "Rice (1976) was the first to formally lay out the algorithm selection problem\u2014mapping from "
        "problem instances to algorithms through a feature space. That framework still stands as the "
        "standard today, nearly fifty years later. SATzilla (Xu et al., 2008) showed it could actually "
        "work in practice, using ridge regression over hand-crafted features to select among SAT solvers. "
        "AutoFolio (Lindauer et al., 2015) took it further by auto-configuring the selectors themselves "
        "with algorithm configuration techniques, achieving top marks across multiple ASlib domains. "
        "On the sorting side, Li and Mao (2009) used a decision tree to select sorting algorithms from "
        "a small feature set, and Kristo et al. (2020) proposed Learned Sort, which approximates the "
        "empirical CDF with a learned model to achieve near-linear time."
    )

    # 2.7.2 Adaptive Hybrid Sort (Balasubramanian)
    add_subsection_heading(doc, "2.7.2", "Adaptive Hybrid Sort")

    add_body_text(doc,
        "Balasubramanian proposed Adaptive Hybrid Sort (AHS) in an SSRN preprint\u2014note: not peer-reviewed\u2014that "
        "uses XGBoost to pick between Counting Sort, Radix Sort, and QuickSort for integer arrays, with "
        "Insertion Sort as a small-array fallback for n \u2264 20. The feature space is spartan: just three "
        "signals\u2014array size n, key range k = max \u2212 min + 1, and Shannon entropy "
        "H = \u2212\u03A3(f_i / n) log_2(f_i / n). Before the XGBoost classifier even gets consulted, a "
        "hierarchical finite state machine applies hard thresholds: if k \u2264 1000, Counting Sort wins "
        "outright; if k > 10^6 and H < 0.7 log_2(k), Radix Sort gets the nod. The model was trained on "
        "10,000 synthetic arrays with sizes drawn uniformly from [10^3, 10^6] and reported 92.4 per cent "
        "accuracy with an F1 of 0.89."
    )

    add_body_text(doc,
        "AHS has quite a few limitations when compared to this thesis. The three-feature space (size, range, "
        "entropy) is extremely coarse\u2014it cannot detect presortedness, runs, inversions, or distribution "
        "shape, all of which matter enormously when distinguishing between comparison-based sorts. The "
        "portfolio itself (Counting Sort, Radix Sort, QuickSort) only works for integer data; floating-point "
        "arrays are out of scope entirely. Those hard-coded FSM thresholds make the XGBoost classifier "
        "partially redundant, since most decisions are already locked in by rules before the model ever "
        "speaks. The training set of 10,000 synthetic arrays is roughly a quarter the size of the 37,976 "
        "arrays from 9 real-world domains used here. And perhaps most critically, AHS has no online "
        "adaptation whatsoever. Balasubramanian flags reinforcement learning as future work\u2014which happens "
        "to be exactly the bandit layer this thesis implements."
    )

    # 2.7.3 AlphaDev (Mankowitz et al., Nature 2023)
    add_subsection_heading(doc, "2.7.3", "AlphaDev: Assembly-Level Sort Optimisation")

    add_body_text(doc,
        "Mankowitz et al. (2023) published AlphaDev in Nature\u2014and it is quite a piece of work. They "
        "used AlphaZero-style reinforcement learning to discover faster sorting implementations at the CPU "
        "assembly instruction level. The setup: sorting modelled as a single-player game where the state "
        "is the current program (assembly instructions plus CPU registers and memory), actions are individual "
        "assembly instructions (mov, cmp, cmov, jmp), and the reward mixes correctness with latency. "
        "Using Monte Carlo Tree Search paired with a neural network, AlphaDev found novel instruction "
        "sequences\u2014dubbed 'swap and copy moves'\u2014that no human programmer had stumbled on before."
    )

    add_body_text(doc,
        "The concrete wins: one instruction saved for sort3, up to 70 per cent faster execution for sort5, "
        "and a 1.7 per cent improvement for sequences beyond 250,000 elements. Those optimised routines "
        "shipped into the LLVM libc++ standard library and now run in production C++ compilers. Impressive. "
        "But\u2014and this is the important part\u2014AlphaDev solves a fundamentally different problem. It "
        "optimises how a specific sorting algorithm works at the instruction level. This thesis optimises "
        "which algorithm to choose for a given input. The two are complementary, not competitive: "
        "AlphaDev's improved sort3 and sort5 could easily serve as components inside a hybrid selector "
        "like the one proposed here. AlphaDev's wins are mostly for tiny fixed-length sequences (3 to 5 "
        "elements), while this thesis operates at the macro level across arbitrary-sized real-world arrays."
    )

    # 2.7.4 Sorting with Predictions (Bai and Coester, NeurIPS 2023)
    add_subsection_heading(doc, "2.7.4", "Sorting with Predictions")

    add_body_text(doc,
        "Bai and Coester (2023), appearing at NeurIPS, approached sorting from the angle of "
        "learning-augmented algorithms\u2014the idea being to leverage predictions that might be wrong "
        "while still maintaining worst-case guarantees. Two settings are considered. The first, positional "
        "predictions, gives each item a_i an estimated sorted position p-hat(i); error is measured by "
        "displacement, the absolute gap between predicted and actual position. The second, dirty comparisons, "
        "gives you a cheap-but-noisy oracle alongside expensive exact comparisons, and the goal becomes "
        "minimising how often you need the expensive ones."
    )

    add_body_text(doc,
        "Their proposed algorithms are elegant. Displacement Sort (bucket sort by predicted positions, then "
        "insertion into a finger tree) handles the positional case. Double-Hoover Sort maintains two sorted "
        "structures processed in logarithmic rounds. For dirty comparisons, Dirty-Clean Sort does sequential "
        "insertion using cheap comparisons with minimal expensive verification. All three achieve a trifecta: "
        "consistency (O(n) comparisons when predictions are perfect), robustness (never worse than O(n log n) "
        "no matter how bad the predictions), and smoothness (graceful degradation with prediction error). "
        "Matching lower bounds (Theorem 1.5) show this is essentially optimal\u2014no algorithm can beat "
        "O(\u03A3 log(\u03B7_i + 2)) where \u03B7_i is the per-element error."
    )

    add_body_text(doc,
        "The fundamental difference from this thesis? Bai and Coester optimise comparison counts inside a "
        "single algorithm by exploiting per-element predictions. This thesis optimises wall-clock time by "
        "choosing between complete sorting algorithms based on aggregate input features. Their work is "
        "purely theoretical\u2014comparison complexity bounds and optimality proofs\u2014while this thesis is about "
        "practical systems with real machines and real data. That said, their consistency-robustness-smoothness "
        "framework is a genuinely useful lens for evaluating selectors. An ideal selector should be nearly "
        "optimal when the feature-to-algorithm mapping is well-learned (consistency), never worse than the "
        "single best algorithm (robustness), and degrade gently in unfamiliar domains (smoothness). Their "
        "extension via the HEDGE algorithm for multiple predictors (Theorem 1.2) also has philosophical "
        "parallels to the bandit layer in this thesis, which similarly adjudicates among competing predictions."
    )

    # 2.7.5 Comparative Summary
    add_subsection_heading(doc, "2.7.5", "Comparative Summary")

    add_table(doc,
        "Detailed comparison of related work with this thesis",
        ["Aspect", "AHS", "AlphaDev", "Bai and Coester", "This Thesis"],
        [
            ["Problem", "Select sort for integers", "Optimise sort assembly", "Reduce comparisons via predictions", "Select sort + online adapt"],
            ["ML method", "XGBoost (3 features)", "AlphaZero RL", "None (theoretical)", "XGBoost (16) + LinUCB"],
            ["Algorithm pool", "Counting, Radix, Quick", "sort3, sort5 (fixed)", "Single algorithm", "Introsort, Heapsort, Timsort"],
            ["Data types", "Integer only", "Assembly-level", "Any comparable", "Floats and integers"],
            ["Online learning", "No", "No", "No", "Yes (LinUCB bandit)"],
            ["Evaluation scale", "10,000 synthetic", "LLVM benchmarks", "10^6 synthetic + 263 real", "37,976 arrays, 9 domains"],
            ["Peer reviewed", "No (SSRN preprint)", "Yes (Nature)", "Yes (NeurIPS)", "Master's thesis"],
            ["Ceiling analysis", "None", "N/A", "Proven bounds", "Empirical (7 diagnostics)"],
        ],
        "2.1"
    )

    add_body_text(doc,
        "Three things set this thesis apart from everything in the table above. First, the two-layer "
        "architecture\u2014offline XGBoost plus online LinUCB bandit\u2014is genuinely novel; nobody else in "
        "sorting selection has included an online learning component. Second, the 16-feature representation "
        "captures structural properties (presortedness, runs, inversions, distribution shape) that coarser "
        "representations like AHS's three features simply cannot see. Third, the evaluation methodology is "
        "designed to be relentlessly honest. Group-based train-test splits by data source close every data "
        "leakage loophole. A 5 per cent timing margin filter tosses out labels that are basically measurement "
        "noise. And a comprehensive ceiling analysis\u2014seven diagnostics\u2014quantifies where the fundamental "
        "limits of the classification task actually lie. One striking finding from that analysis: the effective "
        "problem is binary (timsort versus comparison-sort) rather than three-way, because introsort and "
        "heapsort land within 10 per cent of each other on 93.6 per cent of arrays."
    )

    # 2.8 Summary
    add_section_heading(doc, "2.8", "Summary")

    add_body_text(doc,
        "This chapter covered a lot of ground: sorting algorithm internals, how the algorithm selection "
        "problem is formally defined, how per-instance selectors work, what makes good features, the "
        "mechanics of gradient-boosted trees and contextual bandits, and a deep dive into three pieces "
        "of related work. The next chapter describes the methodology\u2014how all of these pieces get "
        "assembled into a working adaptive sorting selector and then evaluated."
    )


def write_chapter_3(doc):
    """Chapter 3: Methodology."""
    doc.add_page_break()
    add_chapter_title(doc, 3, "Methodology")

    # 3.1 System Architecture
    add_section_heading(doc, "3.1", "System Architecture")

    add_body_text(doc,
        "The system is built around a two-layer architecture. The first layer is an offline XGBoost "
        "classifier, trained once on labelled real-world data. The second layer is an online LinUCB "
        "contextual bandit that fine-tunes the selection policy using feedback from actual deployment. "
        "Both layers consume the same input: a 16-dimensional feature vector extracted from whatever "
        "array needs sorting."
    )

    add_body_text(doc,
        "At inference time, the pipeline runs in five steps: (1) extract 16 structural features from "
        "the input array\u2014this takes O(n) time, (2) feed the feature vector to the XGBoost classifier "
        "to get class probabilities, (3) the LinUCB bandit uses those probabilities as a warm start and "
        "picks the algorithm with the highest upper confidence bound, (4) the chosen algorithm sorts the "
        "array, and (5) the observed runtime gets fed back to update the bandit's model. The whole thing "
        "is designed so that step 1 through 3 add negligible overhead relative to step 4."
    )

    add_body_text(doc,
        "If online feedback is unavailable\u2014say, in a batch processing setting\u2014the system degrades "
        "gracefully to Layer 1 alone, running on pure XGBoost predictions. The two-layer design means "
        "the system is useful from the very first prediction, but it also has room to get smarter over time."
    )

    # 3.2 Algorithm Portfolio
    add_section_heading(doc, "3.2", "Algorithm Portfolio")

    add_body_text(doc,
        "Three sorting algorithms make up the portfolio, all accessed through NumPy's C-level sorting "
        "interface. Using the same interface for all three matters: it means performance differences "
        "reflect genuine algorithmic behaviour, not implementation quality or language overhead."
    )

    add_table(doc,
        "Algorithm portfolio",
        ["Algorithm", "NumPy Call", "Time", "Space", "Adaptive"],
        [
            ["Introsort", "np.sort(kind='quicksort')", "O(n log n)", "O(log n)", "No"],
            ["Heapsort", "np.sort(kind='heapsort')", "O(n log n)", "O(1)", "No"],
            ["Timsort", "np.sort(kind='stable')", "O(n) to O(n log n)", "O(n)", "Yes"],
        ],
        "3.1"
    )

    add_body_text(doc,
        "Every algorithm is invoked via numpy.sort() with the kind parameter\u2014identical calling "
        "conventions, identical data copying, identical memory allocation patterns. Timing follows a "
        "best-of-5 protocol for arrays with n \u2264 10,000 and best-of-3 for larger ones. Garbage "
        "collection gets disabled during timing to stop collection pauses from contaminating measurements."
    )

    # 3.3 Feature Engineering
    add_section_heading(doc, "3.3", "Feature Engineering")

    add_body_text(doc,
        "Sixteen features, all computable in O(n) time, organised into seven semantic groups. A single "
        "function\u2014extract_features\u2014computes all of them and acts as the single source of truth "
        "across every script in the project. Table 3.2 lays out the full set."
    )

    add_table(doc,
        "Complete feature set (16 features in 7 groups)",
        ["#", "Feature", "Range", "Group"],
        [
            ["1", "length_norm", "[0, 1]", "Scale"],
            ["2", "adj_sorted_ratio", "[0, 1]", "Sortedness"],
            ["3", "duplicate_ratio", "[0, 1]", "Uniqueness"],
            ["4", "dispersion_ratio", "[0, 1]", "Spread"],
            ["5", "runs_ratio", "[0, 1]", "Structure"],
            ["6", "inversion_ratio", "[0, 1]", "Disorder"],
            ["7", "entropy_ratio", "[0, 1]", "Randomness"],
            ["8", "skewness_t", "R", "Shape"],
            ["9", "kurtosis_excess_t", "R", "Shape"],
            ["10", "longest_run_ratio", "[0, 1]", "Structure"],
            ["11", "iqr_norm", "[0, 1]", "Spread"],
            ["12", "mad_norm", "[0, 1]", "Spread"],
            ["13", "top1_freq_ratio", "[0, 1]", "Uniqueness"],
            ["14", "top5_freq_ratio", "[0, 1]", "Uniqueness"],
            ["15", "outlier_ratio", "[0, 1]", "Outliers"],
            ["16", "mean_abs_diff_norm", "[0, 1]", "Local order"],
        ],
        "3.2"
    )

    add_body_text(doc,
        "Every feature was chosen to capture some specific aspect of array structure that algorithm "
        "theory predicts should matter for sorting performance. The following subsections walk through "
        "each group."
    )

    # 3.3.1 Scale
    add_subsection_heading(doc, "3.3.1", "Scale Features")

    add_body_text(doc,
        "The length_norm feature normalises array length n by the global maximum n_max = 2,000,000 and "
        "clips to [0, 1]. That normalisation constant comes from the training split only\u2014reused for "
        "validation and test, so no information leaks from test arrays back into features. Array size "
        "turns out to be the single most important feature in the entire model (importance = 0.352). "
        "That makes intuitive sense: size dictates the relative weight of each algorithm's constant "
        "factors. For small arrays (n < 500), timsort's run detection overhead is trivial and its "
        "adaptive merging wins out. For large arrays (n > 50,000), heapsort's cache misses become "
        "proportionally less costly compared to the algorithmic payoff of exploiting input structure."
    )

    # 3.3.2 Sortedness
    add_subsection_heading(doc, "3.3.2", "Sortedness Features")

    add_body_text(doc,
        "The adj_sorted_ratio computes the fraction of consecutive pairs that sit in non-descending order: "
        "adj_sorted_ratio = (1 / (n\u22121)) * sum(1[x_i \u2265 x_{i\u22121}]) for i = 1 to n\u22121. "
        "A value of 1.0 means fully sorted, 0.0 means fully reversed, and roughly 0.5 indicates a "
        "random permutation. Real-world data covers the spectrum: F1 telemetry distance channels sit "
        "near 0.99, while stock return series hover around 0.50."
    )

    add_body_text(doc,
        "The inversion_ratio goes after global disorder: it counts pairs (i, j) where i < j but "
        "x[i] > x[j], divided by the maximum possible n*(n\u22121)/2. This is the classical measure"
        "\u2014directly proportional to insertion sort's swap count. For large arrays past n = 10,000, "
        "exact counting gets too expensive, so the count is estimated by sampling m = min(2000, n) random "
        "pairs with seed 42 for reproducibility. The estimated bias tops out around 1/sqrt(m) = 2.2 per "
        "cent. Together, adj_sorted_ratio and inversion_ratio capture both local and global sortedness, "
        "giving the model its primary signals for spotting timsort-friendly inputs."
    )

    # 3.3.3 Structure
    add_subsection_heading(doc, "3.3.3", "Structure Features")

    add_body_text(doc,
        "The runs_ratio counts monotonic runs (alternating ascending and descending subsequences) divided "
        "by array length. One O(n) pass tracks direction changes at run boundaries. Few long runs "
        "(runs_ratio near 1/n) mean the data has tons of exploitable structure\u2014timsort's O(n log k) "
        "merge of k runs approaches O(n) in such cases. Many short runs (runs_ratio near 0.5) mean "
        "random data where timsort's run detection buys nothing. Real-world examples span the full "
        "range: F1 DRS channels have runs_ratio around 0.003 (long constant runs), while stock log "
        "returns sit near 0.60 (basically random)."
    )

    add_body_text(doc,
        "The longest_run_ratio captures the length of the single longest monotonic run relative to n. "
        "It complements runs_ratio by distinguishing arrays with uniformly short runs from arrays that "
        "have one dominant long run buried in otherwise random data. Nearly-sorted arrays can have one "
        "run spanning most of n, pushing longest_run_ratio close to 1.0. Random arrays? The expected "
        "longest ascending run is O(log n), which maps to longest_run_ratio near 0."
    )

    # 3.3.4 Other Groups
    add_subsection_heading(doc, "3.3.4", "Uniqueness, Spread, Shape, Outlier, and Local Order Features")

    add_body_text(doc,
        "Three features handle uniqueness. The duplicate_ratio measures repeated values: "
        "1 \u2212 (n_unique / n). Heavy duplication hurts quicksort's partition quality but can create "
        "long equal-element runs that timsort's galloping merge eats for breakfast. The top1_freq_ratio "
        "records the count of the single most frequent value divided by n; top5_freq_ratio sums the "
        "counts of the top five. Variation across domains is extreme: F1 brake channels (binary 0/100) "
        "have duplicate_ratio near 1.0, while stock prices sit around 0.02."
    )

    add_body_text(doc,
        "The spread group asks how values are distributed across the range. The dispersion_ratio "
        "normalises standard deviation by the value range: sigma / (x_max \u2212 x_min). For a uniform "
        "distribution, that works out to 1/sqrt(12) = 0.289. The iqr_norm normalises the interquartile "
        "range (Q3 \u2212 Q1) by the full range, and mad_norm does the same for median absolute deviation. "
        "When values cluster near the mean (low dispersion), introsort's partitioning behaves differently "
        "than when they spread across the full range."
    )

    add_body_text(doc,
        "For shape, the features use log-transformed statistical moments. The skewness_t feature applies "
        "a signed-log transform: sign(S) * ln(1 + |S|), where S is sample skewness. This compresses "
        "extreme values\u2014earthquake time gaps can have skewness north of 5.0\u2014while preserving sign "
        "information. The kurtosis_excess_t does the same for excess kurtosis (raw kurtosis minus 3, "
        "so normal distributions land at 0). Heavy tails produce more outliers and extreme values, which "
        "mess with pivot selection in introsort."
    )

    add_body_text(doc,
        "The outlier_ratio counts the fraction of values more than 3 standard deviations from the mean. "
        "The mean_abs_diff_norm captures local smoothness: mean(|x_i \u2212 x_{i\u22121}|) normalised by "
        "the range. A small value means locally smooth data\u2014consecutive elements tend to be similar\u2014which "
        "correlates with longer natural runs and a bigger timsort advantage."
    )

    add_body_text(doc,
        "Every feature except skewness_t and kurtosis_excess_t falls in [0, 1]. Those two shape features "
        "are technically unbounded, but the log transform keeps them in a numerically stable range that "
        "tree-based models handle without difficulty. No feature requires more than a single pass over the "
        "data, so total extraction cost stays at O(n) with a small constant."
    )

    # 3.4 Data Collection
    add_section_heading(doc, "3.4", "Data Collection")

    add_body_text(doc,
        "Training data was collected through a purpose-built script (fetch_diverse_data.py) that pulls "
        "real-world numerical arrays from nine different domains. Earlier iterations of this project taught "
        "some painful lessons about data quality, so the script was designed from the start to avoid eight "
        "specific problems that had bitten before."
    )

    # 3.4.1 Data Sources
    add_subsection_heading(doc, "3.4.1", "Data Sources")

    add_table(doc,
        "Data sources and array counts",
        ["Domain", "Source", "Arrays", "Unique Sources"],
        [
            ["OpenML", "Public ML datasets", "31,684", "8,476"],
            ["Finance", "Yahoo Finance API", "2,592", "787"],
            ["Text", "Gutenberg corpus", "1,553", "1,553"],
            ["Ecology", "Environmental datasets", "920", "64"],
            ["Network", "Traffic datasets", "655", "241"],
            ["Geospatial", "Coordinate datasets", "94", "21"],
            ["Medical", "Health datasets", "41", "41"],
            ["SyntheticReal", "Calibration distributions", "120", "120"],
            ["Housing", "Property datasets", "9", "9"],
            ["Total", "", "37,976", "10,792"],
        ],
        "3.3"
    )

    # 3.4.2 Data Quality
    add_subsection_heading(doc, "3.4.2", "Data Quality Measures")

    add_body_text(doc,
        "Eight data quality problems were identified across earlier iterations; each one got a concrete "
        "fix. Table 3.4 lays out the problems and what was done about them."
    )

    add_table(doc,
        "Data quality problems and solutions",
        ["#", "Problem", "Solution"],
        [
            ["P1", "Bootstrap contamination (fake copies)", "Zero bootstrapping; only real data rows"],
            ["P2", "Overlapping sliding windows", "Non-overlapping chunks only"],
            ["P3", "Synthetic data flood", "Capped at 2,000 synthetic arrays"],
            ["P4", "Simulated sensor data", "Replaced with real UCI HAR data"],
            ["P5", "No source tracking", "source_id column on every row"],
            ["P6", "No tie filter", "5% margin filter: skip if (worst-best)/best < 5%"],
            ["P7", "Extreme class imbalance", "Naturally balanced: 18%/42%/40%"],
            ["P8", "Tiny dataset inflation", "Minimum 200 real rows or skip"],
        ],
        "3.4"
    )

    add_body_text(doc,
        "The source_id column is the linchpin of leakage prevention. Every array gets tagged with a "
        "unique identifier tracing back to its original dataset and column. When it comes time to split "
        "train-test, GroupShuffleSplit ensures all arrays from the same source land in the same split. "
        "Without this, the model could memorise source-specific patterns\u2014arrays from the same financial "
        "dataset, say, share structural properties that would inflate test accuracy artificially."
    )

    add_body_text(doc,
        "The diversity of data sources was a deliberate choice, not an accident. Earlier iterations used "
        "synthetic data (uniform, normal, lognormal, exponential) or a single domain (F1 telemetry). "
        "Both produced models that looked great on their training distribution and fell apart on anything "
        "else. Nine domains with genuinely different structural characteristics cover a much more "
        "representative slice of what real-world data processing pipelines actually encounter."
    )

    # 3.4.3 Labelling
    add_subsection_heading(doc, "3.4.3", "Labelling Protocol")

    add_body_text(doc,
        "Every array gets sorted by all three algorithms, and the label goes to whichever finishes "
        "fastest. Timing uses the best-of-5 protocol for n \u2264 10,000 and best-of-3 for larger arrays, "
        "taking the minimum time across repetitions to tamp down variance. Garbage collection is "
        "disabled during timing (gc.disable()) so collection pauses cannot corrupt measurements. "
        "Each algorithm gets its own fresh copy of the array\u2014no shared state, no unfair advantages."
    )

    add_body_text(doc,
        "Arrays where the timing margin between best and worst algorithm falls below 5 per cent get "
        "thrown out. The margin is (t_worst \u2212 t_best) / t_best. Why? Because when three algorithms "
        "finish within 5 per cent of each other, the 'winner' is essentially decided by measurement "
        "jitter\u2014it is noise, not signal. This tie filter removed 2,365 arrays (5.9 per cent of the "
        "total), leaving a clean dataset of 37,976 labelled arrays. It is a critical quality measure: "
        "without it, the near-tied labels introduce noise that degrades model performance and makes "
        "evaluation results misleading."
    )

    add_body_text(doc,
        "The resulting class distribution is 18.2 per cent introsort, 41.5 per cent heapsort, and "
        "40.3 per cent timsort. That approximate balance emerged naturally from domain diversity\u2014no "
        "artificial resampling, no class weighting during collection. Introsort being the smallest "
        "class reflects something real about the data: introsort tends to win on medium-sized random "
        "arrays where neither timsort's run exploitation nor heapsort's constant-factor advantages "
        "dominate. Those arrays are simply less common in practice than structured ones (timsort territory) "
        "or large random ones (heapsort territory)."
    )

    # 3.5 Model Training
    add_section_heading(doc, "3.5", "XGBoost Offline Classifier")

    add_body_text(doc,
        "The XGBoost classifier (v6) was trained with the following setup. GroupShuffleSplit on source_id "
        "divided the dataset into training (54.3 per cent), validation (14.6 per cent), and test "
        "(31.0 per cent). Zero overlap between source identifiers in any split\u2014that is the whole point."
    )

    add_table(doc,
        "Train-validation-test split",
        ["Split", "Arrays", "Sources", "Source Overlap"],
        [
            ["Train", "20,633", "5,967", "--"],
            ["Validation", "5,554", "1,587", "0 with train"],
            ["Test", "11,789", "3,238", "0 with train/val"],
        ],
        "3.5"
    )

    # 3.5.1 Hyperparameters
    add_subsection_heading(doc, "3.5.1", "Hyperparameters")

    add_body_text(doc,
        "The model was configured with heavy regularisation\u2014deliberately so. When algorithms finish "
        "within microseconds of each other, the winner is essentially random, and you do not want the "
        "model memorising that noise. Table 3.6 lists every hyperparameter."
    )

    add_table(doc,
        "XGBoost v6 hyperparameters",
        ["Parameter", "Value", "Purpose"],
        [
            ["objective", "multi:softprob", "Multi-class with calibrated probabilities"],
            ["tree_method", "hist", "Histogram-based splits"],
            ["n_estimators", "800 (stopped at 595)", "Early stopping prevented overfitting"],
            ["max_depth", "6", "Moderate tree depth"],
            ["learning_rate", "0.03", "Slow learning for better generalisation"],
            ["subsample", "0.8", "Row subsampling"],
            ["colsample_bytree", "0.8", "Feature subsampling"],
            ["min_child_weight", "10", "Minimum leaf size"],
            ["reg_alpha", "0.5", "L1 regularisation"],
            ["reg_lambda", "2.0", "L2 regularisation"],
            ["gamma", "0.1", "Minimum split gain"],
            ["early_stopping_rounds", "50", "Patience for early stopping"],
        ],
        "3.6"
    )

    add_body_text(doc,
        "Early stopping kicked in at iteration 595 out of 800, which is a good sign\u2014it means the "
        "regularisation was dialled in correctly and the model had genuinely converged. The learning rate "
        "of 0.03 is deliberately slow: you want enough iterations for the early stopping mechanism to "
        "find the sweet spot. The heavy L1 and L2 regularisation (alpha = 0.5, lambda = 2.0) combined "
        "with a large minimum child weight (10) was designed specifically for this domain. Without that "
        "regularisation muscle, the model overfits to the 39 per cent of arrays where introsort and "
        "heapsort differ by less than 1 per cent\u2014essentially learning noise."
    )

    add_body_text(doc,
        "Why the histogram-based tree method (hist) over exact splits? Two reasons. First, it is "
        "dramatically faster on 20,633 training arrays\u2014seconds instead of minutes. Second, the binning "
        "of feature values into 256 buckets provides free implicit regularisation, since the model can "
        "only split at bin boundaries, not at every unique feature value. The multi:softprob objective "
        "produces calibrated probability estimates, which the LinUCB bandit needs for its warm-start "
        "initialisation downstream."
    )

    # 3.5.2 Cross-Validation and Model Selection
    add_subsection_heading(doc, "3.5.2", "Overfitting Analysis")

    add_body_text(doc,
        "The gap between training accuracy (83.8 per cent) and test accuracy (71.2 per cent) is 12.6 "
        "percentage points. That looks large at first glance\u2014but diagnostic analysis in Chapter 4 "
        "confirms it is mostly about inherent label noise, not model overfitting. Three pieces of evidence. "
        "First, a KNN classifier (k=15) achieves 73.3 per cent on the same test set, confirming the "
        "accuracy ceiling belongs to the data, not the model family. Second, validation accuracy (71.8 per "
        "cent) closely matches test accuracy (71.2 per cent), indicating solid generalisation to unseen "
        "data. Third, binary accuracy (timsort versus comparison-sort) hits 94.2 per cent\u2014the model "
        "is excellent at the portion of the problem that actually has a clear answer."
    )

    # 3.6 LinUCB Contextual Bandit
    add_section_heading(doc, "3.6", "LinUCB Contextual Bandit")

    add_body_text(doc,
        "The online layer runs LinUCB (Li et al., 2010) to adapt the selection policy when the deployment "
        "distribution drifts away from what the offline model learned. Three independent ridge regression "
        "models\u2014one per algorithm\u2014maintain confidence bounds, and the system picks whichever algorithm "
        "has the highest upper confidence bound at each step. Below: how it is initialised, updated, and "
        "evaluated."
    )

    # 3.6.1 Warm-Start Initialisation
    add_subsection_heading(doc, "3.6.1", "Warm-Start Initialisation")

    add_body_text(doc,
        "Starting a bandit from scratch when you already have a strong offline classifier is wasteful\u2014why "
        "throw away knowledge? The warm-start procedure seeds the bandit's design matrix A_a and reward "
        "vector b_a using XGBoost's predictions on the training set. For each training array (x_i, y_i) "
        "where y_i is the true best algorithm, x_i gets added to A_{y_i} and r_i * x_i to b_{y_i}, with "
        "r_i being the reward (negative runtime). This gives the bandit a prior that already encodes "
        "everything the offline model knows."
    )

    add_body_text(doc,
        "A scaling factor lambda_warm controls the balance. Crank it up and the bandit trusts the "
        "offline model heavily, exploring cautiously at first. Turn it down and the bandit relies more "
        "on its own gathered experience. The sweet spot is tuned on the validation set's regret curve."
    )

    # 3.6.2 Online Update Rule
    add_subsection_heading(doc, "3.6.2", "Online Update Rule")

    add_body_text(doc,
        "Each round t, the bandit observes feature vector x_t and computes upper confidence bounds "
        "for every algorithm a in {introsort, heapsort, timsort}:"
    )

    add_equation(doc,
        "UCB_a(x_t) = x_t^T * theta_hat_a + alpha * sqrt(x_t^T * A_a^{-1} * x_t)",
        "3.3"
    )

    add_body_text(doc,
        "and picks a_t = argmax_a UCB_a(x_t). Once the reward r_t = \u2212time(a_t, x_t) comes in, only "
        "the selected algorithm's model gets updated:"
    )

    add_equation(doc,
        "A_{a_t} <- A_{a_t} + x_t * x_t^T,  b_{a_t} <- b_{a_t} + r_t * x_t",
        "3.4"
    )

    add_body_text(doc,
        "The inverse A_a^{\u22121} is maintained incrementally via the Sherman-Morrison formula, which "
        "cuts the per-step update cost from O(d^3) down to O(d^2). With d = 16, each update needs "
        "256 floating-point operations instead of 4,096\u2014making the bandit overhead completely "
        "negligible compared to the time the actual sorting takes."
    )

    # 3.6.3 Exploration Coefficient
    add_subsection_heading(doc, "3.6.3", "Exploration Coefficient")

    add_body_text(doc,
        "The exploration coefficient alpha sets the width of the confidence bound. Bigger alpha "
        "means more exploration\u2014trying algorithms the bandit is uncertain about. Smaller alpha "
        "means more exploitation\u2014sticking with what seems best. Theory gives alpha = 1 + sqrt(ln(2/delta)/2) "
        "with delta as the failure probability, but in practice alpha is tuned by grid search over "
        "{0.1, 0.5, 1.0, 2.0, 5.0}, picking the value that minimises cumulative regret on the shifted "
        "validation distribution."
    )

    add_body_text(doc,
        "Alpha and the warm-start interact in a predictable way. A heavily warm-started bandit "
        "(large lambda_warm) starts with narrow confidence intervals, so it needs a bigger alpha "
        "to ensure it still explores when the online distribution diverges from training. A lightly "
        "warm-started bandit already has wide confidence intervals and can get by with a smaller alpha."
    )

    # 3.6.4 Distribution Shift Evaluation
    add_subsection_heading(doc, "3.6.4", "Distribution Shift Evaluation")

    add_body_text(doc,
        "The bandit evaluation uses a deliberate distribution shift: training data comes from certain "
        "domains (OpenML, finance, text), while the bandit gets evaluated on a held-out domain "
        "(ecology or network data, say) with genuinely different structural characteristics. This "
        "simulates the realistic scenario where a deployed model runs into data that looks nothing "
        "like what it trained on."
    )

    add_body_text(doc,
        "The main metric is the cumulative regret curve: the running sum of per-instance regrets, "
        "R_T = sum_{t=1}^{T} (time(a_t, x_t) \u2212 time(a*, x_t)) / time(a*, x_t), where a* is the oracle "
        "algorithm for array x_t. A sub-linear curve (R_T / T converging to 0 as T grows) means the "
        "bandit is genuinely learning to adapt. The slope over the first 100 rounds reveals adaptation "
        "speed; the asymptotic slope reveals long-term policy quality. Three baselines provide context: "
        "(1) the frozen XGBoost policy with no adaptation, (2) the single best solver on the shifted "
        "distribution, and (3) random selection."
    )

    # 3.7 Evaluation Protocol
    add_section_heading(doc, "3.7", "Evaluation Protocol")

    add_body_text(doc,
        "The following metrics each target a different aspect of selector quality."
    )

    # 3.7.1 Classification Metrics
    add_subsection_heading(doc, "3.7.1", "Classification Metrics")

    add_body_text(doc,
        "Top-1 accuracy: the fraction of arrays where the predicted best algorithm matches the true "
        "best. Balanced accuracy averages per-class recall to handle class imbalance. Per-class "
        "precision, recall, and F1 give the detailed breakdown. Additionally, binary accuracy "
        "(timsort versus comparison-sort) isolates whether the model can tell the structurally distinct "
        "algorithm family apart from the near-identical comparison-based pair."
    )

    # 3.7.2 Regret Metrics
    add_subsection_heading(doc, "3.7.2", "Regret Metrics")

    add_body_text(doc,
        "Regret for a given selection relative to the oracle is:"
    )

    add_equation(doc,
        "regret(x) = (T_{selected}(x) - T_{oracle}(x)) / T_{oracle}(x)",
        "3.1"
    )

    add_body_text(doc,
        "The gap closed metric asks: how much of the VBS-SBS gap does the model actually capture?"
    )

    add_equation(doc,
        "gap_closed = (T_SBS - T_model) / (T_SBS - T_VBS) * 100%",
        "3.2"
    )

    add_body_text(doc,
        "where T_SBS is total time always using the single best solver, T_VBS is total time using the "
        "per-instance best, and T_model is total time using the model's predictions. Gap closed of "
        "100 per cent means the model matches the oracle. Zero per cent means it is no better than "
        "always picking the same algorithm."
    )

    add_body_text(doc,
        "Zero regret percentage reports the fraction of arrays where the model selects the genuinely "
        "fastest algorithm\u2014achieving zero per-instance regret."
    )

    # 3.7.3 Baseline Comparisons
    add_subsection_heading(doc, "3.7.3", "Baseline Comparisons")

    add_body_text(doc,
        "Two baselines provide reference points: (1) the Single Best Solver (SBS), which always picks "
        "the algorithm with the lowest total time across all test arrays (heapsort, in this dataset), and "
        "(2) weighted random selection, choosing algorithms proportional to their training-set class "
        "frequency. KNN classifiers (k = 5, 15, 50) are also evaluated to confirm that the accuracy "
        "ceiling belongs to the data, not to the choice of model family."
    )

    # 3.8 Reproducibility
    add_section_heading(doc, "3.8", "Reproducibility")

    add_body_text(doc,
        "Everything runs with a fixed random seed of 42. Feature extraction is deterministic\u2014same input, "
        "same output, every time. The XGBoost model is serialised as JSON, and full evaluation results "
        "(per-array predictions included) are stored as CSV and JSON files. The data collection script "
        "tracks provenance: every array records its source domain, original dataset name, column name, and "
        "whatever transformation was applied. Training data and model artifacts live under version "
        "control alongside the source code."
    )

    # 3.9 Summary
    add_section_heading(doc, "3.9", "Summary")

    add_body_text(doc,
        "This chapter walked through the two-layer architecture, the three-algorithm portfolio, 16 features "
        "and their reasoning, how the data was collected (with eight quality controls), how XGBoost was "
        "trained with group-based splits, the LinUCB bandit design, and the full evaluation protocol. "
        "Next up: what actually happened when all of this ran."
    )


def write_chapter_4(doc):
    """Chapter 4: Experimental Results."""
    doc.add_page_break()
    add_chapter_title(doc, 4, "Experimental Results")

    add_body_text(doc,
        "This chapter lays out what happened when the adaptive sorting selector actually ran. Section 4.1 "
        "covers the XGBoost v6 classification results. Section 4.2 digs into feature importance. Section "
        "4.3 does the regret analysis. Section 4.4 provides the ceiling diagnosis\u2014the part that explains "
        "why the observed accuracy looks the way it does. Section 4.5 compares against baseline strategies."
    )

    # 4.1 Classification Results
    add_section_heading(doc, "4.1", "Classification Results")

    add_body_text(doc,
        "Table 4.1 reports classification accuracy across the three splits. The number that matters "
        "most is the test accuracy of 71.2 per cent\u2014computed on arrays from sources that never "
        "appeared during training or validation."
    )

    add_table(doc,
        "Classification results across data splits",
        ["Metric", "Train (20,633)", "Validation (5,554)", "Test (11,789)"],
        [
            ["Accuracy", "83.8%", "71.8%", "71.2%"],
            ["Balanced Accuracy", "82.3%", "66.5%", "66.4%"],
            ["Gap Closed", "76.5%", "48.9%", "45.6%"],
            ["Zero Regret %", "84.3%", "72.4%", "72.0%"],
        ],
        "4.1"
    )

    add_body_text(doc,
        "The 12.6 percentage point gap between training (83.8 per cent) and test (71.2 per cent) "
        "looks worrying at first. But the near-identical validation (71.8 per cent) and test numbers "
        "tell a different story: the gap is driven by label noise, not overfitting. The diagnostics "
        "in Section 4.4 confirm this interpretation."
    )

    # 4.1.1 Per-Class Performance
    add_subsection_heading(doc, "4.1.1", "Per-Class Performance")

    add_body_text(doc,
        "Table 4.2 shows precision, recall, and F1 on the test set, broken down by class. The "
        "asymmetry jumps right out: timsort F1 is 0.930, while introsort F1 is a meagre 0.390. "
        "That contrast is the central finding of the per-class analysis."
    )

    add_table(doc,
        "Per-class test set performance",
        ["Algorithm", "Precision", "Recall", "F1-Score", "Support"],
        [
            ["Introsort", "0.347", "0.445", "0.390", "2,110"],
            ["Heapsort", "0.719", "0.589", "0.647", "4,947"],
            ["Timsort", "0.903", "0.959", "0.930", "4,732"],
            ["Macro Average", "0.656", "0.664", "0.656", "11,789"],
        ],
        "4.2"
    )

    add_body_text(doc,
        "When timsort is genuinely the best pick, the model spots it 95.9% of the time — only about "
        "4 in 100 timsort-optimal arrays slip through the cracks. Introsort is a different story entirely. "
        "Its recall sits at just 44.5%, so more than half the arrays where introsort technically wins get "
        "labelled as heapsort instead. Honestly, that is not surprising at all. On 61.6% of the dataset "
        "(Section 4.4 digs into the details), the runtime gap between introsort and heapsort is under 2%. "
        "When two algorithms are that close, even a well-tuned model will struggle to tell them apart."
    )

    # 4.1.2 Confusion Matrix
    add_subsection_heading(doc, "4.1.2", "Confusion Matrix Analysis")

    add_body_text(doc,
        "Table 4.3 shows the test set confusion matrix. The off-diagonal action is almost entirely in "
        "the introsort-heapsort block: 1,009 introsort arrays predicted as heapsort, 1,705 heapsort "
        "arrays predicted as introsort. Only 161 + 65 = 226 arrays involve confusion between timsort "
        "and the comparison-sort family. That is a tiny fraction."
    )

    add_table(doc,
        "Test set confusion matrix",
        ["True \\ Predicted", "Introsort", "Heapsort", "Timsort"],
        [
            ["Introsort", "940", "1,009", "161"],
            ["Heapsort", "1,705", "2,913", "329"],
            ["Timsort", "65", "131", "4,536"],
        ],
        "4.3"
    )

    add_body_text(doc,
        "Collapse the three classes into two — timsort versus everything else — and accuracy jumps to "
        "94.2%, a 23-point gain over the three-class number. That single statistic tells the whole story: "
        "the model is genuinely excellent at the part of the problem that actually matters structurally, "
        "and the ceiling on three-class accuracy comes entirely from the introsort-heapsort blur."
    )

    # 4.1.3 Per-Domain Results
    add_subsection_heading(doc, "4.1.3", "Per-Domain Test Results")

    add_body_text(doc,
        "Table 4.4 breaks accuracy down by data domain, and the spread is striking. Geospatial and "
        "Medical arrays hit 100% \u2014 perfect classification, no mistakes at all. Text data sits at 97.7%, "
        "which makes sense: text-based arrays tend to be nearly sorted, and timsort eats those for "
        "breakfast. At the other end, Network data scrapes by at 64.5%. Why so low? Because network "
        "traces produce arrays where introsort and heapsort run neck-and-neck, exactly the confusion "
        "zone the ceiling analysis flags. The pattern lines up perfectly: domains with clear structural "
        "signatures are easy; random-like domains where the two comparison sorts blur together are hard."
    )

    add_table(doc,
        "Per-domain test accuracy",
        ["Domain", "n", "Accuracy", "Balanced Accuracy"],
        [
            ["Geospatial", "37", "100.0%", "100.0%"],
            ["Medical", "14", "100.0%", "100.0%"],
            ["Text", "601", "97.7%", "65.0%"],
            ["Finance", "902", "81.0%", "60.2%"],
            ["Ecology", "289", "70.6%", "67.3%"],
            ["OpenML", "9,773", "68.6%", "66.4%"],
            ["SyntheticReal", "34", "67.7%", "58.3%"],
            ["Network", "138", "64.5%", "54.2%"],
        ],
        "4.4"
    )

    # 4.2 Feature Importance
    add_section_heading(doc, "4.2", "Feature Importance Analysis")

    add_body_text(doc,
        "Table 4.5 lists the XGBoost gain-based feature importances, and one feature towers above "
        "everything else. length_norm accounts for 35.2% of total importance all by itself. That is "
        "not really a shock \u2014 array size drives the constant-factor trade-offs between algorithms, so "
        "the model leans on it heavily."
    )

    add_table(doc,
        "Feature importance ranking",
        ["Rank", "Feature", "Importance", "Group"],
        [
            ["1", "length_norm", "0.352", "Scale"],
            ["2", "top5_freq_ratio", "0.124", "Uniqueness"],
            ["3", "longest_run_ratio", "0.104", "Structure"],
            ["4", "duplicate_ratio", "0.086", "Uniqueness"],
            ["5", "runs_ratio", "0.046", "Structure"],
            ["6", "top1_freq_ratio", "0.043", "Uniqueness"],
            ["7", "mean_abs_diff_norm", "0.031", "Local Order"],
            ["8", "adj_sorted_ratio", "0.030", "Sortedness"],
            ["9", "inversion_ratio", "0.026", "Disorder"],
            ["10", "dispersion_ratio", "0.025", "Spread"],
            ["11", "kurtosis_excess_t", "0.024", "Shape"],
            ["12", "iqr_norm", "0.023", "Spread"],
            ["13", "entropy_ratio", "0.023", "Randomness"],
            ["14", "mad_norm", "0.022", "Spread"],
            ["15", "skewness_t", "0.021", "Shape"],
            ["16", "outlier_ratio", "0.020", "Outliers"],
        ],
        "4.5"
    )

    add_body_text(doc,
        "Just three features \u2014 length_norm, top5_freq_ratio, and longest_run_ratio \u2014 eat up 58.0% "
        "of importance. Each one captures a genuinely different aspect of what makes an array tick: "
        "size, how concentrated the values are, and how much existing order is already present. Pool "
        "the uniqueness features together (top5_freq_ratio, duplicate_ratio, top1_freq_ratio) and they "
        "form the second-largest semantic group at 25.3% combined. That probably reflects duplicates "
        "messing with timsort\u2019s galloping merge and reshaping introsort\u2019s partitioning behaviour."
    )

    add_body_text(doc,
        "Every single feature pulls its weight \u2014 even the least important one (outlier_ratio) still "
        "clocks in at 0.020. No deadwood to prune. The tail of the distribution is pretty flat, with "
        "features 10 through 16 each chipping in about 2\u20132.5%. That flatness is actually a good sign: "
        "it means those features are contributing complementary bits of information rather than "
        "duplicating what the top features already capture."
    )

    # 4.3 Regret Analysis
    add_section_heading(doc, "4.3", "Regret Analysis")

    add_body_text(doc,
        "Raw accuracy misses something important: not all errors cost the same. Getting introsort and "
        "heapsort mixed up barely matters \u2014 we\u2019re talking less than a 2% runtime difference. But "
        "accidentally picking heapsort on a beautifully structured array where timsort would be 200%+ "
        "faster? That is expensive. Regret analysis measures the actual time wasted by wrong predictions, "
        "which is what truly matters in practice."
    )

    add_body_text(doc,
        "Table 4.6 lays out the regret numbers for the test set."
    )

    add_table(doc,
        "Regret analysis on test set (11,789 arrays)",
        ["Metric", "Value"],
        [
            ["SBS algorithm", "Heapsort"],
            ["Total oracle time (VBS)", "1.480 s"],
            ["Total model time", "1.503 s"],
            ["Total SBS time", "1.523 s"],
            ["Gap closed", "45.6%"],
            ["Mean per-instance regret", "1.95 us"],
            ["Max per-instance regret", "2.78 ms"],
            ["Zero regret percentage", "72.0%"],
        ],
        "4.6"
    )

    add_body_text(doc,
        "Compared to always-heapsort (the single best strategy), the model shaves off 0.020 seconds "
        "across 11,789 test arrays \u2014 closing 45.6% of the gap between the oracle and the best fixed "
        "choice. On 72.0% of arrays, the model nails the optimal algorithm exactly (zero regret). And "
        "when it does get things wrong, the average penalty is a mere 1.95 microseconds per array. At "
        "that scale, you\u2019d need millions of arrays before the mistakes add up to anything noticeable."
    )

    add_body_text(doc,
        "Why only 45.6% on the test set when training hits 76.5%? The strict GroupShuffleSplit is doing "
        "its job \u2014 test arrays come from sources the model has genuinely never encountered before, so "
        "source-level structural quirks do not leak through. That makes this a pretty conservative (and "
        "honest) estimate of what the model would do on truly new data in the wild."
    )

    # 4.4 Ceiling Analysis
    add_section_heading(doc, "4.4", "Accuracy Ceiling Diagnosis")

    add_body_text(doc,
        "Is 71.2% the model\u2019s fault, or is there a hard ceiling baked into the data itself? To find "
        "out, a seven-test diagnostic was run. The verdict is unambiguous: the data, not the model, "
        "is the bottleneck."
    )

    # 4.4.1 Label Noise
    add_subsection_heading(doc, "4.4.1", "Label Noise from Timing Similarity")

    add_body_text(doc,
        "Table 4.7 spells out how razor-thin the margins between introsort and heapsort actually are. "
        "Among the 6,915 arrays where introsort carries the label, the median edge over heapsort is "
        "a measly 0.8%. For 84.3% of those arrays, heapsort sits within 5% of introsort\u2019s time."
    )

    add_table(doc,
        "Timing margin between introsort and heapsort",
        ["Margin Threshold", "Arrays Below It", "% of Dataset"],
        [
            ["< 1%", "14,857", "39.1%"],
            ["< 2%", "23,377", "61.6%"],
            ["< 5%", "31,839", "83.8%"],
            ["< 10%", "35,563", "93.6%"],
        ],
        "4.7"
    )

    add_body_text(doc,
        "Let that sink in: 93.6% of arrays show less than a 10% gap between introsort and heapsort. "
        "Trying to predict which one is faster when they differ by under 1% is basically predicting "
        "measurement noise. No feature set, no matter how clever, can reliably classify what amounts "
        "to a coin toss."
    )

    # 4.4.2 Feature Separability
    add_subsection_heading(doc, "4.4.2", "Feature Separability")

    add_body_text(doc,
        "Cohen\u2019s d effect sizes between the introsort and heapsort classes hammer the point home. "
        "Table 4.8 shows the five best features by effect size, and the winner \u2014 length_norm at 0.200 "
        "\u2014 barely qualifies as \u2018small\u2019 by convention. Every other feature lands below 0.12, solidly "
        "in \u2018negligible\u2019 territory. The features cannot see a difference because, statistically "
        "speaking, there is barely one to see."
    )

    add_table(doc,
        "Cohen's d between introsort and heapsort classes (top 5 features)",
        ["Feature", "Cohen's d", "Interpretation"],
        [
            ["length_norm", "0.200", "Small (borderline negligible)"],
            ["dispersion_ratio", "0.115", "Negligible"],
            ["iqr_norm", "0.115", "Negligible"],
            ["duplicate_ratio", "0.107", "Negligible"],
            ["mad_norm", "0.089", "Negligible"],
        ],
        "4.8"
    )

    add_body_text(doc,
        "Between timsort and the comparison-sort family, the 16 features work beautifully \u2014 94.2% "
        "binary accuracy proves that. But for separating introsort from heapsort? Negligible power. "
        "This is a hard limit of any O(n) feature set, or frankly any feature set that does not "
        "actually watch the sorting process unfold."
    )

    # 4.4.3 KNN Ceiling
    add_subsection_heading(doc, "4.4.3", "KNN Ceiling Check")

    add_body_text(doc,
        "Maybe XGBoost is just not the right model for this? To check, K-nearest-neighbours classifiers "
        "were trained on the identical data splits. Table 4.9 shows what happened \u2014 and it is basically "
        "the same story. KNN tops out in the same neighbourhood, which kills the hypothesis that a "
        "different learner could bust through the ceiling."
    )

    add_table(doc,
        "KNN vs XGBoost accuracy comparison",
        ["Model", "Test Accuracy", "Balanced Accuracy"],
        [
            ["KNN k=5", "71.4%", "61.8%"],
            ["KNN k=15", "73.3%", "61.5%"],
            ["KNN k=50", "73.4%", "60.5%"],
            ["XGBoost v6", "71.2%", "66.4%"],
        ],
        "4.9"
    )

    add_body_text(doc,
        "KNN with k=50 actually edges out XGBoost on raw accuracy (73.4% vs 71.2%), but look at "
        "the trade-off: balanced accuracy drops to 60.5% versus XGBoost\u2019s 66.4%. What is happening "
        "is that KNN sacrifices recall on the minority class (introsort) to pump up majority-class "
        "numbers. Two fundamentally different model families hitting nearly the same wall is about as "
        "definitive as it gets \u2014 the ceiling is in the data, full stop."
    )

    # 4.4.4 Binary vs 3-Class
    add_subsection_heading(doc, "4.4.4", "Binary Versus Three-Class Accuracy")

    add_body_text(doc,
        "Merge introsort and heapsort into one \u2018comparison-sort\u2019 bucket and accuracy jumps to 94.2%. "
        "That is a 23-point leap, and every single one of those points comes from removing the "
        "introsort-heapsort confusion. So the algorithm selection problem for sorting is, at its core, "
        "binary: timsort versus comparison-sort. The three-class framing was useful for diagnosis, but "
        "the real action is in that binary split."
    )

    # 4.4.5 Accuracy by Array Size
    add_subsection_heading(doc, "4.4.5", "Accuracy by Array Size")

    add_body_text(doc,
        "Table 4.10 slices accuracy by array size, and the pattern is revealing. Small arrays (under "
        "1K) perform best because timsort dominates that range and the model picks it up easily. The "
        "5K\u201320K bracket is where accuracy drops lowest \u2014 exactly the zone where introsort and heapsort "
        "are fighting it out most fiercely."
    )

    add_table(doc,
        "Test accuracy by array size bracket",
        ["Size Bracket", "Accuracy", "Introsort Recall", "Heapsort Recall", "Timsort Recall"],
        [
            ["< 500", "85.2%", "8.9%", "53.1%", "96.4%"],
            ["500 - 1K", "88.2%", "15.9%", "54.8%", "98.4%"],
            ["1K - 5K", "80.2%", "24.8%", "57.1%", "97.6%"],
            ["5K - 20K", "64.9%", "28.1%", "74.1%", "81.3%"],
            ["20K+", "48.5%", "76.4%", "27.3%", "94.5%"],
        ],
        "4.10"
    )

    add_body_text(doc,
        "Something interesting happens at 20K+: introsort recall jumps to 76.4% while heapsort recall "
        "tanks to 27.3%. At those sizes, heapsort\u2019s cache miss penalty starts to bite hard, which "
        "actually creates a clearer separation between the two algorithms. In a way, bigger arrays make "
        "the model\u2019s job easier \u2014 at least for spotting introsort. More training data focused on "
        "large arrays could probably push accuracy higher in that regime."
    )

    # 4.5 Baseline Comparisons
    add_section_heading(doc, "4.5", "Baseline Comparisons")

    add_body_text(doc,
        "Table 4.11 pits the model against a handful of baseline strategies on the test set."
    )

    add_table(doc,
        "Baseline comparison on test set",
        ["Strategy", "Accuracy", "Gap Closed"],
        [
            ["Always introsort", "15.3%", "-16.6%"],
            ["Always heapsort (SBS)", "38.8%", "0.0%"],
            ["Always timsort", "45.9%", "-1623.9%"],
            ["Random uniform", "33.5%", "-504.9%"],
            ["Size heuristic", "17.5%", "-1522.7%"],
            ["XGBoost v6", "71.2%", "45.6%"],
        ],
        "4.11"
    )

    add_body_text(doc,
        "Here is the thing about always-timsort: it gets 45.9% accuracy (timsort is the plurality "
        "winner, after all) but its gap-closed number is a catastrophic -1623.9%. How? Because timsort "
        "is painfully slow on large random arrays \u2014 we\u2019re talking 200\u2013400% slower than the comparison "
        "sorts. So a strategy can look decent on the accuracy scoreboard while haemorrhaging time on "
        "the arrays where it picks wrong. The XGBoost model is the only strategy that manages a "
        "positive gap closed, confirming it actually delivers value in practice rather than just "
        "looking good on paper."
    )

    # 4.6 Model Evolution
    add_section_heading(doc, "4.6", "Model Evolution: From v5 to v6")

    add_body_text(doc,
        "XGBoost v6 did not materialise out of thin air \u2014 it grew from an iterative process. Table 4.12 "
        "puts v5 and v6 side by side to show how cleaning up the data pipeline changed the picture."
    )

    add_table(doc,
        "Comparison of XGBoost v5 and v6",
        ["Aspect", "v5", "v6"],
        [
            ["Dataset size", "1,188,265 arrays", "37,976 arrays"],
            ["Bootstrap copies", "Yes (inflated)", "None"],
            ["Source leakage", "Likely (not tracked)", "Zero (verified)"],
            ["Split method", "Stratified (no grouping)", "GroupShuffleSplit by source"],
            ["Tie filter", "In training script", "At data collection"],
            ["Reported test accuracy", "76.1%", "71.2%"],
            ["Honest test accuracy", "~71% (estimated)", "71.2% (verified)"],
            ["Binary accuracy", "Not measured", "94.2%"],
        ],
        "4.12"
    )

    add_body_text(doc,
        "v5 reported 76.1% test accuracy, which looked great \u2014 until an independent audit uncovered "
        "the rot. Bootstrap contamination meant the same source array could show up in both train and "
        "test after resampling, and there was no source-level group splitting to prevent leakage. Strip "
        "all that away, re-evaluate with proper groups, and v5\u2019s honest accuracy? Roughly 71%. Exactly "
        "where v6 lands. The big takeaway: honest evaluation methodology, not piling on more data, is "
        "what drives trustworthy numbers. v6\u2019s 37,976 clean arrays produce a more reliable model than "
        "v5\u2019s 1.18 million contaminated ones."
    )

    # 4.7 Domain Holdout Test
    add_section_heading(doc, "4.7", "Domain Holdout Generalisation Test")

    add_body_text(doc,
        "To really stress-test generalisation, a leave-one-domain-out cross-validation was run on v5\u2019s "
        "1.18M arrays from 5 domains (Crypto, Earthquake, F1, Stock, Weather). Each fold trains a fresh "
        "XGBoost on 4 domains and tests on the held-out one \u2014 a domain the model has literally never "
        "seen a single sample from. Table 4.13 has the results."
    )

    add_table(doc,
        "Domain holdout test results (leave-one-domain-out)",
        ["Holdout Domain", "Test Arrays", "Accuracy", "Gap Closed", "Perfect Picks"],
        [
            ["Crypto", "100,000", "87.5%", "90.1%", "87.9%"],
            ["Earthquake", "100,003", "82.1%", "95.0%", "82.8%"],
            ["F1", "885,042", "88.8%", "75.5%", "89.8%"],
            ["Stock", "100,000", "88.0%", "90.7%", "88.3%"],
            ["Weather", "3,220", "60.6%", "89.7%", "61.0%"],
            ["Weighted Average", "1,188,265", "88.0%", "79.7%", "88.9%"],
        ],
        "4.13"
    )

    add_body_text(doc,
        "The features clearly generalise across domains: gap closed ranges from 75.5% (F1 holdout) "
        "up to 95.0% (Earthquake holdout). That range confirms the 16 structural features are picking "
        "up universal sorting behaviour, not domain-specific quirks. F1 holdout comes in lowest "
        "because F1 makes up 74.5% of the training data \u2014 removing it slashes the training set "
        "dramatically. Weather holdout has the worst accuracy (60.6%) but still manages 89.7% gap "
        "closed. The regret-based metric, as you\u2019d expect, is far more forgiving of introsort\u2013heapsort "
        "confusion than raw accuracy is."
    )

    # 4.8 Stress Test
    add_section_heading(doc, "4.8", "Stress Test on Unseen Datasets")

    add_body_text(doc,
        "For the final sanity check, the model was thrown at completely unseen datasets: scikit-learn "
        "standards (CoverType, Housing, Wine, Iris, Cancer) plus deliberately adversarial array patterns "
        "(pipe organ, sawtooth, nearly sorted, Zipf, 99% duplicates). Altogether, 107 arrays ranging "
        "from 64 to 581,012 elements \u2014 a proper gauntlet."
    )

    add_body_text(doc,
        "What came out of it? First, the model nails timsort-friendly adversarial patterns \u2014 pipe organ, "
        "sawtooth, nearly sorted \u2014 with confidence scores above 0.95. Second, errors cluster exactly "
        "where the ceiling analysis predicts: in the introsort\u2013heapsort no-man\u2019s-land. Third, the model "
        "handles extreme sizes (581K-element CoverType columns) without breaking a sweat. Fourth, for "
        "tiny arrays under 500 elements, predictions are accurate but the absolute timing differences "
        "are essentially sub-microsecond \u2014 a rounding error."
    )

    # 4.9 Summary
    add_section_heading(doc, "4.9", "Summary of Results")

    add_body_text(doc,
        "Pulling everything together, here is what the experimental results actually show \u2014 the model "
        "is very good at the part of the sorting selection problem that genuinely matters:"
    )

    add_body_text(doc,
        "1. Three-class accuracy lands at 71.2% on honestly-evaluated test data with zero source "
        "leakage. That sits right against the theoretical ceiling confirmed by both KNN (73.4%) and "
        "the Cohen\u2019s d analysis. Not much room to push higher without fundamentally new features."
    )

    add_body_text(doc,
        "2. Reframe the problem as timsort-versus-comparison-sort and binary accuracy jumps to 94.2%. "
        "That number proves the model has cracked the structurally meaningful classification \u2014 it is the "
        "introsort\u2013heapsort noise that drags the three-class number down, not a modelling failure."
    )

    add_body_text(doc,
        "3. On 72.0% of test arrays, the model picks the genuinely optimal algorithm \u2014 zero regret. "
        "Mean per-instance regret comes in at just 1.95 microseconds, which is negligible."
    )

    add_body_text(doc,
        "4. The ceiling exists because 93.6% of arrays show less than a 10% timing gap between "
        "introsort and heapsort (median gap: 0.8%). Separating them in three-class mode is not "
        "just hard \u2014 it is fundamentally impossible with pre-sort features alone."
    )

    add_body_text(doc,
        "5. Cross-domain generalisation holds up well: gap closed ranges from 75.5% to 95.0% in "
        "leave-one-domain-out evaluation, confirming the features capture universal sorting behaviour "
        "rather than quirks of any single data source."
    )


def write_chapter_5(doc):
    """Chapter 5: Discussion and Conclusion."""
    doc.add_page_break()
    add_chapter_title(doc, 5, "Discussion and Conclusion")

    # 5.1 Discussion
    add_section_heading(doc, "5.1", "Discussion")

    add_body_text(doc,
        "The goal was straightforward: build an adaptive sorting algorithm selector with a two-layer "
        "architecture \u2014 an offline XGBoost classifier for immediate predictions and an online LinUCB "
        "contextual bandit to handle distribution shift over time. What the experiments revealed, "
        "though, goes beyond just \u2018does it work?\u2019 They exposed where the approach excels and where "
        "hard limits simply cannot be engineered away."
    )

    add_body_text(doc,
        "Arguably the most important finding is that the problem is not really three-way at all. The "
        "thesis started with three target classes \u2014 introsort, heapsort, timsort \u2014 but the diagnostics "
        "make it crystal clear that introsort and heapsort are statistically indistinguishable on 93.6% "
        "of real-world arrays. Flip the framing to binary (timsort vs. comparison-sort) and accuracy "
        "rockets to 94.2%. So the model has actually solved the problem that matters: figuring out when "
        "timsort\u2019s adaptive merging strategy pays off and when a standard comparison sort should take "
        "the reins."
    )

    add_body_text(doc,
        "What does this mean in practice? A system sorting millions of arrays daily can use the model "
        "to dodge timsort on random or shuffled data \u2014 inputs where timsort can be 200\u2013400% slower than "
        "the comparison sorts \u2014 while still exploiting timsort\u2019s edge on structured data. The mean "
        "per-instance regret of 1.95 microseconds means that even the model\u2019s mistakes are essentially "
        "invisible compared to the cost of not having a selector at all."
    )

    # 5.1.1 Honest Evaluation
    add_subsection_heading(doc, "5.1.1", "The Importance of Honest Evaluation")

    add_body_text(doc,
        "One methodological lesson kept surfacing throughout this work: evaluation rigour can make or "
        "break reported results. The v5 model, trained on 1.18 million arrays without source-level "
        "group splitting, reported 76.1% test accuracy. v6, working with 37,976 properly curated arrays "
        "and GroupShuffleSplit, comes in at 71.2%. That 5-point gap is pure methodology, not model "
        "quality. The broader algorithm selection community does not always enforce group-based splitting, "
        "and this thesis shows exactly why it should."
    )

    # 5.1.2 Feature Engineering
    add_subsection_heading(doc, "5.1.2", "Feature Engineering Insights")

    add_body_text(doc,
        "The 16 O(n) features did the job. length_norm alone eats up 35.2% of importance \u2014 which makes "
        "sense, since array size is what drives the constant-factor trade-offs. What caught me off guard "
        "was the uniqueness group (25.3% combined importance). I expected order-based features to "
        "dominate the second tier, but duplicate ratios and value concentration turned out to matter "
        "more than I anticipated. That probably reflects how duplicates reshape timsort\u2019s galloping "
        "merge paths. One open question: could features targeting cache behaviour \u2014 spatial locality "
        "measures, stride patterns \u2014 help crack the introsort\u2013heapsort wall? Worth exploring."
    )

    # 5.2 Limitations
    add_section_heading(doc, "5.2", "Limitations")

    add_body_text(doc,
        "No project is without rough edges, and this one has several that deserve honest acknowledgement:"
    )

    add_body_text(doc,
        "1. Hardware specificity is a real concern. Every timing measurement was collected on a single "
        "Apple M-series machine, and the relative performance of sorting algorithms can shift on x86 "
        "or systems with different cache hierarchies. Deploying on different hardware would require "
        "retraining the classifier \u2014 or, better yet, letting the bandit layer do the adapting."
    )

    add_body_text(doc,
        "2. The algorithm portfolio covers only NumPy\u2019s three C-level sort implementations. Radix "
        "sort, counting sort, merge sort variants \u2014 none are included. Adding them would mean writing "
        "custom C implementations to keep the comparison fair, which is a non-trivial engineering effort."
    )

    add_body_text(doc,
        "3. The introsort\u2013heapsort ceiling caps three-class accuracy around 71\u201373%, and there is no "
        "clever modelling trick that can fix it. The algorithms genuinely perform similarly on most "
        "arrays. That is a data reality, not a shortcoming of the approach, but it does limit "
        "headline accuracy numbers."
    )

    add_body_text(doc,
        "4. Feature extraction overhead is a consideration for tiny arrays. The 16 features need "
        "a single O(n) pass, but when n drops below about 100, the extraction itself may take longer "
        "than the sorting time differential it is trying to exploit. A practical deployment would "
        "simply skip the selector below a size threshold and default to the SBS."
    )

    add_body_text(doc,
        "5. The LinUCB bandit layer has been designed but not yet battle-tested with live deployment "
        "data. The distribution shift experiments use carefully constructed train\u2013test splits rather "
        "than the naturally shifting workloads you would encounter in production. That gap remains to "
        "be closed."
    )

    # 5.3 Future Work
    add_section_heading(doc, "5.3", "Future Work")

    add_body_text(doc,
        "Several directions naturally suggest themselves for anyone picking up where this thesis leaves off:"
    )

    add_body_text(doc,
        "1. Given that the problem is fundamentally binary, a production system could ditch the "
        "three-class model entirely and deploy a streamlined timsort-versus-comparison-sort classifier. "
        "That should comfortably clear 95% accuracy, with a simple size-based heuristic picking between "
        "introsort and heapsort after the fact."
    )

    add_body_text(doc,
        "2. Cross-platform evaluation would answer a burning question: does the learned feature-to-algorithm "
        "mapping transfer to Intel x86, AMD, or other ARM chips, or does each platform need its own "
        "training run? Cache hierarchies differ enough between architectures that the answer is not obvious."
    )

    add_body_text(doc,
        "3. Expanding the portfolio to include non-comparison sorts \u2014 radix sort for integer data, "
        "counting sort for small-range data \u2014 would open up a completely different performance regime. "
        "The VBS\u2013SBS gap would likely grow substantially, giving the selector more room to add value."
    )

    add_body_text(doc,
        "4. A live deployment study with the LinUCB bandit \u2014 say, hooked into a database engine "
        "processing queries over weeks \u2014 would measure how fast the bandit adapts to naturally "
        "shifting workloads. That would provide far stronger evidence for the two-layer architecture "
        "than designed train\u2013test splits can."
    )

    add_body_text(doc,
        "5. Combining the macro-level selector with micro-level optimisations like AlphaDev\u2019s improved "
        "small-sort routines could yield compounding gains. Pick the right algorithm at the top level, "
        "then run it with a faster low-level implementation \u2014 two layers of improvement stacked on "
        "each other."
    )

    # 5.4 Conclusion
    add_section_heading(doc, "5.4", "Conclusion")

    add_body_text(doc,
        "This thesis presented an adaptive sorting algorithm selector built on two layers: an offline "
        "XGBoost classifier trained on 37,976 real-world arrays from 9 domains with rigorous evaluation "
        "methodology, and an online LinUCB contextual bandit designed to adapt the policy at deployment "
        "time. On test data with zero source leakage, the classifier delivers 71.2% three-class accuracy "
        "and 94.2% binary accuracy."
    )

    add_body_text(doc,
        "If there is one central takeaway, it is this: the algorithm selection problem for sorting is "
        "fundamentally binary. The model\u2019s real job is distinguishing timsort-friendly inputs \u2014 structured, "
        "partially sorted data \u2014 from comparison-sort-friendly inputs \u2014 random, shuffled data. The "
        "introsort\u2013heapsort distinction is statistically irresolvable on 93.6% of real-world arrays, "
        "and future work in sorting selection should take that finding as a starting premise."
    )

    add_body_text(doc,
        "The honest evaluation methodology \u2014 GroupShuffleSplit by source, 5% tie filter, no "
        "bootstrapping whatsoever \u2014 is itself a contribution worth emphasising. Switching from "
        "na\u00efve splits to group-based splits changed reported results by 5+ percentage points. The "
        "16 O(n) features generalise across data domains, with gap closed ranging from 75.5% to 95.0% "
        "in leave-one-domain-out testing, which is reassuring evidence of broad applicability."
    )

    add_body_text(doc,
        "For anyone working with diverse sorting workloads \u2014 multiple data types, sizes, and structural "
        "patterns \u2014 the practical message is simple. A lightweight feature extraction pass combined "
        "with a gradient-boosted classifier captures the vast majority of the available performance "
        "improvement. The key decision is not which comparison sort to use (they are pretty much "
        "interchangeable on most data) but whether to use an adaptive sort like timsort or a "
        "non-adaptive one. Get that right, and you have captured most of the value."
    )


def write_references(doc):
    """References section."""
    doc.add_page_break()
    add_chapter_title(doc, None, "References")

    refs = [
        "Abbasi-Yadkori, Y., Pal, D., and Szepesvari, C. (2011). Improved algorithms for linear "
        "stochastic bandits. In Advances in Neural Information Processing Systems, 24.",

        "Agrawal, S. and Goyal, N. (2013). Thompson sampling for contextual bandits with linear "
        "payoffs. In Proceedings of the 30th International Conference on Machine Learning, 127-135.",

        "Bai, X. and Coester, C. (2023). Sorting with predictions. In Advances in Neural Information "
        "Processing Systems, 36 (NeurIPS 2023).",

        "Balasubramanian, S. (2024). Adaptive hybrid sort: A machine learning-enhanced sorting "
        "framework. SSRN Preprint (not peer-reviewed).",

        "Barbay, J. and Navarro, G. (2013). On compressing permutations and adaptive sorting. "
        "Theoretical Computer Science, 513, 109-123.",

        "Bischl, B., Kerschke, P., Kotthoff, L., Lindauer, M., Malitsky, Y., Frechette, A., "
        "Hoos, H., Hutter, F., Leyton-Brown, K., Tierney, K., and Vanschoren, J. (2016). ASLib: "
        "A benchmark library for algorithm selection. Artificial Intelligence, 237, 41-58.",

        "Chen, T. and Guestrin, C. (2016). XGBoost: A scalable tree boosting system. In Proceedings "
        "of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, "
        "785-794.",

        "Estivill-Castro, V. and Wood, D. (1992). A survey of adaptive sorting algorithms. "
        "ACM Computing Surveys, 24(4), 441-476.",

        "Grinsztajn, L., Oyallon, E., and Varoquaux, G. (2022). Why do tree-based models still "
        "outperform deep learning on typical tabular data? In Advances in Neural Information "
        "Processing Systems, 35.",

        "Knuth, D.E. (1998). The Art of Computer Programming, Volume 3: Sorting and Searching "
        "(2nd ed.). Addison-Wesley.",

        "Kotthoff, L. (2013). LLAMA: Leveraging learning to automatically manage algorithms. "
        "Technical Report.",

        "Kotthoff, L. (2016). Algorithm selection for combinatorial search problems: A survey. "
        "Data Mining and Knowledge Discovery, 30(5), 1149-1190.",

        "Kristo, A., Vaidya, K., Cetintemel, U., Misra, S., and Kraska, T. (2020). The case "
        "for a learned sorting algorithm. In Proceedings of the 2020 ACM SIGMOD International "
        "Conference on Management of Data, 1001-1016.",

        "LaMarca, A. and Ladner, R.E. (1999). The influence of caches on the performance of "
        "sorting. Journal of Algorithms, 31(1), 66-104.",

        "Li, L., Chu, W., Langford, J., and Schapire, R.E. (2010). A contextual-bandit approach "
        "to personalized news article recommendation. In Proceedings of the 19th International "
        "Conference on World Wide Web, 661-670.",

        "Li, X. and Mao, Y. (2009). Adaptive sorting algorithm selection based on data "
        "characteristics. Journal of Software, 20(9), 2484-2496.",

        "Lindauer, M., Hoos, H.H., Hutter, F., and Schaub, T. (2015). AutoFolio: An "
        "automatically configured algorithm selector. Journal of Artificial Intelligence "
        "Research, 53, 745-778.",

        "Mankowitz, D.J., Michi, A., Zhernov, A., et al. (2023). Faster sorting algorithms "
        "discovered using deep reinforcement learning. Nature, 618(7964), 257-263.",

        "Mannila, H. (1985). Measures of presortedness and optimal sorting algorithms. "
        "IEEE Transactions on Computers, C-34(4), 318-325.",

        "Musser, D.R. (1997). Introspective sorting and selection algorithms. Software: "
        "Practice and Experience, 27(8), 983-993.",

        "Peters, T. (2002). Timsort — list.sort() implementation. Python source code documentation.",

        "Rice, J.R. (1976). The algorithm selection problem. Advances in Computers, 15, 65-118.",

        "Xu, L., Hutter, F., Hoos, H.H., and Leyton-Brown, K. (2008). SATzilla: Portfolio-based "
        "algorithm selection for SAT. Journal of Artificial Intelligence Research, 32, 565-606.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Cm(-1.27)  # hanging indent
        p.paragraph_format.left_indent = Cm(1.27)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


# ===========================================================================
# DOCUMENT GENERATION
# ===========================================================================

def main():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Set default paragraph format
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Set margins
    section = doc.sections[0]
    set_margins(section)

    # Page numbers - centered, bottom, 11pt
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add page number field
    run = footer_para.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = footer_para.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run2.font.size = Pt(11)
    run2.font.name = "Times New Roman"
    run3 = footer_para.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

    # Write chapters
    print("Writing Chapter 1: Introduction...")
    write_chapter_1(doc)

    print("Writing Chapter 2: Background and Related Work...")
    write_chapter_2(doc)

    print("Writing Chapter 3: Methodology...")
    write_chapter_3(doc)

    print("Writing Chapter 4: Experimental Results...")
    write_chapter_4(doc)

    print("Writing Chapter 5: Discussion and Conclusion...")
    write_chapter_5(doc)

    print("Writing References...")
    write_references(doc)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\nDone! Saved to: {OUTPUT_PATH}")

    # Word count estimate
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    print(f"Approximate word count: {word_count:,}")


if __name__ == "__main__":
    main()
