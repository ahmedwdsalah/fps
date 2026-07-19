#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import shutil
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import xgboost as xgb
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from feature_extraction import FEATURE_NAMES

ROOT = Path(__file__).resolve().parent.parent
DATA_CSV = ROOT / "data" / "training_dataset.csv"
MODEL_PATH = ROOT / "models" / "xgboost_v5" / "xgb_v5.json"
OUT = ROOT / "results" / "final3_behavior_loop"

ALGORITHMS = ["introsort", "heapsort", "timsort"]
TIME_COLS = {a: f"time_{a}" for a in ALGORITHMS}
FEATURE_GROUPS = {
    "size": ["n_elements", "length_norm"],
    "order": ["adj_sorted_ratio", "runs_ratio", "inversion_ratio", "longest_run_ratio", "mean_abs_diff_norm"],
    "repetition": ["duplicate_ratio", "top1_freq_ratio", "top5_freq_ratio"],
    "distribution": ["entropy_ratio", "dispersion_ratio", "iqr_norm", "mad_norm", "skewness_t", "kurtosis_excess_t", "outlier_ratio"],
}
ALL_FEATURES = ["n_elements"] + FEATURE_NAMES
CLASS_MAP = {0: "heapsort", 1: "introsort", 2: "timsort"}


@dataclass
class Paths:
    tables: Path
    figures: Path
    checkpoints: Path
    report_md: Path
    report_docx: Path
    state_json: Path


def paths() -> Paths:
    return Paths(
        tables=OUT / "tables",
        figures=OUT / "figures",
        checkpoints=OUT / "checkpoints",
        report_md=OUT / "knowledge_report.md",
        report_docx=OUT / "final3_behavior_knowledge_report.docx",
        state_json=OUT / "state.json",
    )


def ensure_dirs(p: Paths) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    p.tables.mkdir(exist_ok=True)
    p.figures.mkdir(exist_ok=True)
    p.checkpoints.mkdir(exist_ok=True)


def save_json(path: Path, data: object) -> None:
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def validate_input(df: pd.DataFrame) -> None:
    required = ["file", "domain", "n_elements", *FEATURE_NAMES, *TIME_COLS.values()]
    missing = [c for c in required if c not in df.columns]
    if missing:
        raise SystemExit(f"missing required columns: {missing}")
    times = df[list(TIME_COLS.values())]
    if not np.isfinite(times.to_numpy()).all():
        raise SystemExit("timing columns contain NaN or infinite values")
    if (times <= 0).any().any():
        raise SystemExit("timing columns contain non-positive values")


def load_data(sample: int | None) -> pd.DataFrame:
    cols = ["file", "domain", "n_elements", *FEATURE_NAMES, *TIME_COLS.values()]
    df = pd.read_csv(DATA_CSV, usecols=cols, nrows=sample)
    validate_input(df)
    times = df[[TIME_COLS[a] for a in ALGORITHMS]].to_numpy()
    df["true_fastest"] = np.array(ALGORITHMS)[times.argmin(axis=1)]
    df["oracle_time"] = times.min(axis=1)
    sorted_times = np.sort(times, axis=1)
    df["runtime_margin_us"] = (sorted_times[:, 1] - sorted_times[:, 0]) * 1e6
    df["oracle_time_us"] = df["oracle_time"] * 1e6
    return df


def add_predictions(df: pd.DataFrame) -> tuple[pd.DataFrame, str]:
    if not MODEL_PATH.exists():
        return df, f"model missing: {MODEL_PATH}"
    model = xgb.XGBClassifier()
    model.load_model(str(MODEL_PATH))
    pred_raw = model.predict(df[FEATURE_NAMES].to_numpy())
    df["predicted"] = [CLASS_MAP[int(x)] for x in pred_raw]
    pred_time = np.zeros(len(df), dtype=np.float64)
    for algo in ALGORITHMS:
        mask = df["predicted"].eq(algo).to_numpy()
        pred_time[mask] = df.loc[mask, TIME_COLS[algo]].to_numpy()
    df["predicted_time"] = pred_time
    df["success"] = df["predicted"].eq(df["true_fastest"])
    df["regret_us"] = (df["predicted_time"] - df["oracle_time"]) * 1e6
    df.loc[np.abs(df["regret_us"]) < 1e-9, "regret_us"] = 0.0
    df["failure_type"] = np.where(df["success"], "success", np.where(df["regret_us"] <= 1.0, "low_regret_failure", "high_regret_failure"))
    return df, "predictions generated from models/xgboost_v5/xgb_v5.json"


def write_stage(p: Paths, stage: str, title: str, lines: list[str]) -> None:
    ck = p.checkpoints / f"{stage}.md"
    body = [f"## {title}", "", *lines, ""]
    ck.write_text("\n".join(body), encoding="utf-8")
    rebuild_reports(p)


def rebuild_reports(p: Paths) -> None:
    order = [
        "01_baseline",
        "02_winner_regions",
        "03_success_indicators",
        "04_failure_indicators",
        "05_case_mining",
        "06_domain_patterns",
        "07_interaction_patterns",
        "08_class_specific_patterns",
        "09_worst_failure_anatomy",
        "10_synthesis",
    ]
    chunks = ["# Final-Three Sorting Behavior Knowledge Report", "", f"Generated: {datetime.now().isoformat(timespec='seconds')}", ""]
    for name in order:
        f = p.checkpoints / f"{name}.md"
        if f.exists():
            chunks.append(f.read_text(encoding="utf-8"))
    p.report_md.write_text("\n".join(chunks), encoding="utf-8")
    build_docx_report(p, chunks)


def build_docx_report(p: Paths, chunks: list[str]) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    title = doc.add_paragraph()
    r = title.add_run("Final-Three Sorting Behavior Knowledge Report")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    for chunk in chunks[3:]:
        for line in chunk.splitlines():
            if line.startswith("## "):
                para = doc.add_paragraph()
                run = para.add_run(line[3:])
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(31, 78, 121)
            elif line.startswith("- "):
                para = doc.add_paragraph(style=None)
                para.paragraph_format.left_indent = Inches(0.18)
                run = para.add_run(clean_docx_text(line[2:]))
                run.font.name = "Times New Roman"
                run.font.size = Pt(10.5)
            elif line.strip():
                para = doc.add_paragraph(clean_docx_text(line))
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    for fig in sorted(p.figures.glob("*.png")):
        doc.add_page_break()
        cap = doc.add_paragraph(fig.stem.replace("_", " ").title())
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        doc.add_picture(str(fig), width=Inches(6.2))
    doc.save(p.report_docx)


def clean_docx_text(text: str) -> str:
    return re.sub(r"`([^`]*)`", r"\1", text)


def pct_table(series: pd.Series) -> pd.DataFrame:
    vc = series.value_counts().reindex(ALGORITHMS, fill_value=0)
    return pd.DataFrame({"algorithm": vc.index, "count": vc.values, "pct": vc.values / len(series) * 100})


def share_frame(df: pd.DataFrame, group_cols: list[str], target: str) -> pd.DataFrame:
    rows = []
    for keys, part in df.groupby(group_cols, observed=False):
        if not isinstance(keys, tuple):
            keys = (keys,)
        vc = part[target].value_counts(normalize=True)
        row = {c: v for c, v in zip(group_cols, keys)}
        row.update({
            "n": int(len(part)),
            "dominant": vc.idxmax(),
            "dominant_pct": float(vc.max() * 100),
            "timsort_pct": float(vc.get("timsort", 0) * 100),
            "introsort_pct": float(vc.get("introsort", 0) * 100),
            "heapsort_pct": float(vc.get("heapsort", 0) * 100),
        })
        if "success" in part:
            row["success_pct"] = float(part["success"].mean() * 100)
            row["mean_regret_us"] = float(part["regret_us"].mean())
            row["p95_regret_us"] = float(part["regret_us"].quantile(0.95))
        row["median_margin_us"] = float(part["runtime_margin_us"].median())
        rows.append(row)
    return pd.DataFrame(rows)


def stage_baseline(df: pd.DataFrame, p: Paths) -> None:
    winner = pct_table(df["true_fastest"])
    winner.to_csv(p.tables / "baseline_winner_share.csv", index=False)
    totals = []
    for algo in ALGORITHMS:
        s = df[TIME_COLS[algo]]
        totals.append({"algorithm": algo, "total_s": s.sum(), "median_us": s.median() * 1e6, "p95_us": s.quantile(0.95) * 1e6})
    totals_df = pd.DataFrame(totals).sort_values("total_s")
    totals_df.to_csv(p.tables / "baseline_runtime_summary.csv", index=False)
    vbs = df["oracle_time"].sum()
    sbs_algo = totals_df.iloc[0]["algorithm"]
    sbs = totals_df.iloc[0]["total_s"]
    gap = 100 * (sbs - vbs) / sbs
    save_json(p.tables / "baseline_summary.json", {"rows": int(len(df)), "sbs_algorithm": sbs_algo, "vbs_total_s": vbs, "sbs_total_s": sbs, "vbs_sbs_gap_pct": gap})
    plot_winner_share(winner, p.figures / "baseline_winner_share.png")
    lines = [
        f"- Rows analyzed: {len(df):,}.",
        f"- SBS algorithm: `{sbs_algo}`.",
        f"- VBS total: {vbs:.6f}s; SBS total: {sbs:.6f}s; VBS/SBS gap: {gap:.2f}%.",
        "- Winner share: " + ", ".join(f"{r.algorithm} {r.pct:.2f}%" for r in winner.itertuples()),
        "- Meaning: this is the raw runtime landscape before explaining success or failure.",
    ]
    write_stage(p, "01_baseline", "Baseline Runtime Landscape", lines)


def bucket_ids(s: pd.Series, q: int = 10) -> pd.Series:
    ranked = s.rank(method="first")
    try:
        return pd.qcut(ranked, q=min(q, max(1, s.nunique())), labels=False, duplicates="drop")
    except ValueError:
        return pd.Series([0] * len(s), index=s.index)


def bucket_label(values: pd.Series) -> str:
    lo = float(values.min())
    hi = float(values.max())
    if lo == hi:
        return f"{lo:.4g}"
    return f"{lo:.4g} to {hi:.4g}"


def region_table(df: pd.DataFrame, features: list[str], target_col: str, prefix: str, p: Paths) -> pd.DataFrame:
    rows = []
    for feat in features:
        b = bucket_ids(df[feat])
        tmp = df.assign(_bucket=b)
        for bucket, part in tmp.groupby("_bucket", observed=False):
            vc = part[target_col].value_counts(normalize=True)
            winner = vc.idxmax()
            rows.append({
                "feature": feat,
                "bucket": bucket_label(part[feat]),
                "n": int(len(part)),
                "dominant": winner,
                "dominant_pct": float(vc.max() * 100),
                "timsort_pct": float(vc.get("timsort", 0) * 100),
                "introsort_pct": float(vc.get("introsort", 0) * 100),
                "heapsort_pct": float(vc.get("heapsort", 0) * 100),
                "median_margin_us": float(part["runtime_margin_us"].median()),
            })
    out = pd.DataFrame(rows).sort_values(["dominant_pct", "n"], ascending=[False, False])
    out.to_csv(p.tables / f"{prefix}.csv", index=False)
    return out


def stage_winner_regions(df: pd.DataFrame, p: Paths) -> None:
    regions = region_table(df, ALL_FEATURES, "true_fastest", "winner_regions_by_feature_bucket", p)
    top = regions[regions["n"] >= max(20, int(len(df) * 0.005))].head(12)
    plot_top_regions(top, p.figures / "winner_region_top_indicators.png", "strongest true-winner regions")
    lines = [
        "- Each structural feature was bucketed into quantile regions.",
        "- For each region, the loop measured which of timsort/introsort/heapsort is the dominant runtime winner.",
        "- Strongest regions found:",
    ]
    lines += [f"- `{r.feature}` {r.bucket}: {r.dominant} dominates with {r.dominant_pct:.1f}% over {r.n:,} arrays; median winner margin {r.median_margin_us:.2f} us." for r in top.itertuples()]
    write_stage(p, "02_winner_regions", "Structural Regions Where Each Algorithm Wins", lines)


def stage_success(df: pd.DataFrame, p: Paths) -> bool:
    if "success" not in df:
        write_stage(p, "03_success_indicators", "Why Correct Predictions Are Correct", ["- Prediction analysis skipped because model predictions were not available."])
        return False
    ok = df[df["success"]].copy()
    tab = region_table(ok, ALL_FEATURES, "true_fastest", "success_indicators_by_feature_bucket", p)
    by_class = ok.groupby("true_fastest").agg(n=("true_fastest", "size"), median_margin_us=("runtime_margin_us", "median"), median_oracle_us=("oracle_time_us", "median")).reset_index()
    by_class.to_csv(p.tables / "success_by_true_class.csv", index=False)
    top = tab[tab["n"] >= max(20, int(len(ok) * 0.005))].head(12)
    plot_top_regions(top, p.figures / "success_top_indicators.png", "strongest success indicators")
    lines = [
        f"- Successful predictions: {len(ok):,} ({100 * len(ok) / len(df):.2f}%).",
        "- Highest success indicators are regions where the true fastest class is structurally clear and the runtime margin is usually not tiny.",
    ]
    lines += [f"- `{r.feature}` {r.bucket}: correct {r.dominant} region, {r.dominant_pct:.1f}% dominance, n={r.n:,}, median margin={r.median_margin_us:.2f} us." for r in top.itertuples()]
    write_stage(p, "03_success_indicators", "Why Correct Predictions Are Correct", lines)
    return True


def stage_failure(df: pd.DataFrame, p: Paths) -> bool:
    if "success" not in df:
        write_stage(p, "04_failure_indicators", "Why Failed Predictions Fail", ["- Failure analysis skipped because model predictions were not available."])
        return False
    fail = df[~df["success"]].copy()
    if fail.empty:
        write_stage(p, "04_failure_indicators", "Why Failed Predictions Fail", ["- No failures found in this run."])
        return True
    fail["pair"] = fail["true_fastest"] + "_to_" + fail["predicted"]
    pair = fail.groupby("pair").agg(n=("pair", "size"), mean_regret_us=("regret_us", "mean"), median_regret_us=("regret_us", "median"), p95_regret_us=("regret_us", lambda x: x.quantile(0.95))).sort_values(["n", "mean_regret_us"], ascending=False).reset_index()
    pair.to_csv(p.tables / "failure_pairs.csv", index=False)
    high = fail[fail["regret_us"] > 1.0]
    low = fail[fail["regret_us"] <= 1.0]
    high_regions = region_table(high if len(high) else fail, ALL_FEATURES, "pair", "high_regret_failure_indicators_by_feature_bucket", p)
    top = high_regions[high_regions["n"] >= max(10, int(max(len(high), 1) * 0.005))].head(12)
    plot_failure_pairs(pair.head(10), p.figures / "failure_pairs_regret.png")
    lines = [
        f"- Failed predictions: {len(fail):,} ({100 * len(fail) / len(df):.2f}%).",
        f"- Low-regret failures <=1 us: {len(low):,}; high-regret failures >1 us: {len(high):,}.",
        "- Main failed class-pairs:",
    ]
    lines += [f"- `{r.pair}`: n={r.n:,}, mean regret={r.mean_regret_us:.2f} us, p95={r.p95_regret_us:.2f} us." for r in pair.head(8).itertuples()]
    lines.append("- Strong high-regret feature regions:")
    lines += [f"- `{r.feature}` {r.bucket}: {r.dominant} dominates failures with {r.dominant_pct:.1f}%, n={r.n:,}." for r in top.itertuples()]
    write_stage(p, "04_failure_indicators", "Why Failed Predictions Fail", lines)
    return True


def stage_cases(df: pd.DataFrame, p: Paths) -> None:
    cases = []
    for algo in ALGORITHMS:
        part = df[df["true_fastest"].eq(algo)]
        if "success" in df:
            part = part[part["success"]]
        if not part.empty:
            cases.append(part.sort_values("runtime_margin_us", ascending=False).head(1))
    if "success" in df:
        fail = df[~df["success"]].copy()
        if not fail.empty:
            cases.append(fail.sort_values("regret_us", ascending=True).head(1))
            cases.append(fail.sort_values("regret_us", ascending=False).head(1))
    out = pd.concat(cases, ignore_index=True) if cases else df.head(0)
    keep = ["file", "domain", "n_elements", "true_fastest", "runtime_margin_us", "oracle_time_us", *TIME_COLS.values()]
    if "predicted" in out:
        keep += ["predicted", "success", "regret_us", "failure_type"]
    keep += FEATURE_NAMES
    out[keep].to_csv(p.tables / "representative_cases.csv", index=False)
    lines = ["- Representative cases exported to `tables/representative_cases.csv`."]
    for r in out.itertuples():
        pred = getattr(r, "predicted", "not_available")
        regret = getattr(r, "regret_us", 0.0)
        lines.append(f"- `{r.domain}` n={r.n_elements:,}: true={r.true_fastest}, pred={pred}, regret={regret:.2f} us, margin={r.runtime_margin_us:.2f} us, file=`{r.file}`.")
    write_stage(p, "05_case_mining", "Representative Examples", lines)


def stage_domain_patterns(df: pd.DataFrame, p: Paths) -> None:
    tab = share_frame(df, ["domain"], "true_fastest").sort_values(["dominant_pct", "n"], ascending=[False, False])
    tab.to_csv(p.tables / "domain_winner_success_patterns.csv", index=False)
    plot_domain_winners(tab, p.figures / "domain_winner_patterns.png")
    lines = [
        "- Domain split tests whether global pattern hides different behavior per source.",
        "- Domain winner and success patterns:",
    ]
    for r in tab.itertuples():
        extra = f", success={r.success_pct:.2f}%, mean regret={r.mean_regret_us:.2f} us" if hasattr(r, "success_pct") else ""
        lines.append(f"- `{r.domain}`: n={r.n:,}, dominant={r.dominant} {r.dominant_pct:.1f}%, tim={r.timsort_pct:.1f}%, intro={r.introsort_pct:.1f}%, heap={r.heapsort_pct:.1f}%{extra}.")
    write_stage(p, "06_domain_patterns", "Domain-Level Behavior Patterns", lines)


def stage_interaction_patterns(df: pd.DataFrame, p: Paths) -> None:
    work = df.copy()
    work["size_bucket"] = bucket_ids(work["n_elements"], q=5)
    work["size_range"] = work.groupby("size_bucket")["n_elements"].transform(bucket_label)
    for feat in ["adj_sorted_ratio", "inversion_ratio", "runs_ratio", "duplicate_ratio", "top5_freq_ratio"]:
        work[f"{feat}_bucket"] = bucket_ids(work[feat], q=5)
        work[f"{feat}_range"] = work.groupby(f"{feat}_bucket")[feat].transform(bucket_label)
    frames = []
    for feat in ["adj_sorted_ratio", "inversion_ratio", "runs_ratio", "duplicate_ratio", "top5_freq_ratio"]:
        tab = share_frame(work, ["size_range", f"{feat}_range"], "true_fastest")
        tab.insert(0, "interaction", f"n_elements x {feat}")
        tab = tab.rename(columns={f"{feat}_range": "feature_range"})
        frames.append(tab)
    out = pd.concat(frames, ignore_index=True).sort_values(["dominant_pct", "n"], ascending=[False, False])
    out.to_csv(p.tables / "size_feature_interaction_patterns.csv", index=False)
    top = out[out["n"] >= max(50, int(len(df) * 0.003))].head(14)
    plot_interaction_patterns(top, p.figures / "size_feature_interaction_patterns.png")
    lines = [
        "- Interaction test checks whether one feature still matters after size changes.",
        "- Strongest size x structure regions:",
    ]
    for r in top.itertuples():
        extra = f", success={r.success_pct:.1f}%, p95 regret={r.p95_regret_us:.2f} us" if hasattr(r, "success_pct") else ""
        lines.append(f"- `{r.interaction}` size {r.size_range}, feature {r.feature_range}: {r.dominant} {r.dominant_pct:.1f}%, n={r.n:,}{extra}.")
    write_stage(p, "07_interaction_patterns", "Size and Structure Interaction Patterns", lines)


def stage_class_specific_patterns(df: pd.DataFrame, p: Paths) -> None:
    rows = []
    for algo in ALGORITHMS:
        part = df[df["true_fastest"].eq(algo)]
        for feat in ALL_FEATURES:
            b = bucket_ids(part[feat])
            tmp = part.assign(_bucket=b)
            for _, sub in tmp.groupby("_bucket", observed=False):
                if len(sub) == 0:
                    continue
                row = {
                    "true_class": algo,
                    "feature": feat,
                    "bucket": bucket_label(sub[feat]),
                    "n": int(len(sub)),
                    "share_of_class_pct": float(len(sub) / len(part) * 100) if len(part) else 0.0,
                    "median_margin_us": float(sub["runtime_margin_us"].median()),
                    "median_oracle_us": float(sub["oracle_time_us"].median()),
                }
                if "success" in sub:
                    row["success_pct"] = float(sub["success"].mean() * 100)
                    row["mean_regret_us"] = float(sub["regret_us"].mean())
                rows.append(row)
    out = pd.DataFrame(rows).sort_values(["true_class", "success_pct" if "success" in df else "share_of_class_pct", "median_margin_us"], ascending=[True, False, False])
    out.to_csv(p.tables / "class_specific_success_patterns.csv", index=False)
    lines = [
        "- Per-class test avoids `timsort` dominance hiding intro/heap signals.",
        "- Best class-specific indicators:",
    ]
    for algo in ALGORITHMS:
        sub = out[(out["true_class"].eq(algo)) & (out["n"] >= max(20, int((df["true_fastest"].eq(algo).sum()) * 0.05)))].head(5)
        for r in sub.itertuples():
            extra = f", success={r.success_pct:.1f}%" if hasattr(r, "success_pct") else ""
            lines.append(f"- `{algo}`: `{r.feature}` {r.bucket}, n={r.n:,}, class share={r.share_of_class_pct:.1f}%, margin={r.median_margin_us:.2f} us{extra}.")
    write_stage(p, "08_class_specific_patterns", "Class-Specific Success Indicators", lines)


def stage_worst_failure_anatomy(df: pd.DataFrame, p: Paths) -> None:
    if "success" not in df:
        write_stage(p, "09_worst_failure_anatomy", "Worst Failure Anatomy", ["- Skipped because predictions were not available."])
        return
    fail = df[~df["success"]].copy()
    if fail.empty:
        write_stage(p, "09_worst_failure_anatomy", "Worst Failure Anatomy", ["- No failures found."])
        return
    fail["pair"] = fail["true_fastest"] + "_to_" + fail["predicted"]
    worst = fail.sort_values("regret_us", ascending=False).head(500)
    worst.to_csv(p.tables / "worst_500_failures.csv", index=False)
    by_pair = worst.groupby("pair").agg(n=("pair", "size"), mean_regret_us=("regret_us", "mean"), median_n=("n_elements", "median"), median_margin_us=("runtime_margin_us", "median")).sort_values(["n", "mean_regret_us"], ascending=False).reset_index()
    by_pair.to_csv(p.tables / "worst_failure_pair_anatomy.csv", index=False)
    feat_summary = []
    for feat in ["n_elements", "adj_sorted_ratio", "inversion_ratio", "runs_ratio", "longest_run_ratio", "duplicate_ratio", "top5_freq_ratio", "entropy_ratio"]:
        feat_summary.append({
            "feature": feat,
            "worst_median": float(worst[feat].median()),
            "all_failure_median": float(fail[feat].median()),
            "all_data_median": float(df[feat].median()),
        })
    feat_df = pd.DataFrame(feat_summary)
    feat_df.to_csv(p.tables / "worst_failure_feature_shift.csv", index=False)
    lines = [
        "- Worst-failure test isolates top 500 errors by regret, not by count.",
        "- Worst failure pairs:",
    ]
    for r in by_pair.head(8).itertuples():
        lines.append(f"- `{r.pair}`: n={r.n:,}, mean regret={r.mean_regret_us:.2f} us, median n={r.median_n:.0f}, median margin={r.median_margin_us:.2f} us.")
    lines.append("- Feature shift in worst failures:")
    for r in feat_df.itertuples():
        lines.append(f"- `{r.feature}`: worst median={r.worst_median:.4g}, all-failure median={r.all_failure_median:.4g}, all-data median={r.all_data_median:.4g}.")
    write_stage(p, "09_worst_failure_anatomy", "Worst Failure Anatomy", lines)


def stage_synthesis(df: pd.DataFrame, p: Paths, prediction_status: str) -> None:
    lines = [
        "- The loop builds a behavior map for the final three algorithms only.",
        "- The main question is which structural indicators make timsort, introsort, or heapsort win.",
        "- Success analysis explains where the model sees the same structure as the timing oracle.",
        "- Failure analysis separates harmless boundary mistakes from expensive mistakes.",
        f"- Prediction status: {prediction_status}.",
        "- Next tests should target the strongest high-regret regions with controlled synthetic probes and optional fresh retiming.",
    ]
    write_stage(p, "10_synthesis", "Final Synthesis and Next Tests", lines)
    save_json(p.state_json, {"finished_at": datetime.now().isoformat(timespec="seconds"), "rows": int(len(df)), "prediction_status": prediction_status})


def plot_winner_share(winner: pd.DataFrame, path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7, 4), dpi=180)
    ax.bar(winner["algorithm"], winner["pct"], color=["#1f77b4", "#888888", "#2ca02c"])
    ax.set_ylabel("winner share (%)")
    ax.set_title("True fastest algorithm share")
    for i, v in enumerate(winner["pct"]):
        ax.text(i, v + 1, f"{v:.1f}%", ha="center", fontweight="bold")
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_top_regions(tab: pd.DataFrame, path: Path, title: str) -> None:
    if tab.empty:
        return
    labels = [f"{r.feature}\n{r.dominant}" for r in tab.itertuples()]
    fig, ax = plt.subplots(figsize=(9, 4.8), dpi=180)
    ax.bar(range(len(tab)), tab["dominant_pct"], color="#df982d")
    ax.set_xticks(range(len(tab)), labels, rotation=35, ha="right")
    ax.set_ylabel("dominant share (%)")
    ax.set_title(title)
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_failure_pairs(tab: pd.DataFrame, path: Path) -> None:
    if tab.empty:
        return
    fig, ax = plt.subplots(figsize=(8, 4.6), dpi=180)
    ax.barh(tab["pair"][::-1], tab["mean_regret_us"][::-1], color="#c89116")
    ax.set_xlabel("mean regret (us)")
    ax.set_title("Failed prediction pairs by mean regret")
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="x", alpha=0.2)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_domain_winners(tab: pd.DataFrame, path: Path) -> None:
    if tab.empty:
        return
    fig, ax = plt.subplots(figsize=(8.5, 4.8), dpi=180)
    x = np.arange(len(tab))
    width = 0.24
    ax.bar(x - width, tab["timsort_pct"], width, label="timsort", color="#2ca02c")
    ax.bar(x, tab["introsort_pct"], width, label="introsort", color="#1f77b4")
    ax.bar(x + width, tab["heapsort_pct"], width, label="heapsort", color="#888888")
    ax.set_xticks(x, tab["domain"], rotation=20, ha="right")
    ax.set_ylabel("true winner share (%)")
    ax.set_title("Winner distribution by domain")
    ax.legend(frameon=False, ncols=3)
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def plot_interaction_patterns(tab: pd.DataFrame, path: Path) -> None:
    if tab.empty:
        return
    labels = [f"{r.interaction.replace('n_elements x ', '')}\n{r.dominant}" for r in tab.itertuples()]
    fig, ax = plt.subplots(figsize=(9.5, 5.2), dpi=180)
    ax.bar(range(len(tab)), tab["dominant_pct"], color="#df982d")
    ax.set_xticks(range(len(tab)), labels, rotation=35, ha="right")
    ax.set_ylabel("dominant winner share (%)")
    ax.set_title("Strong size x structure regions")
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def reset_outputs(p: Paths) -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    ensure_dirs(p)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--sample", type=int, default=None, help="limit rows for dry-run validation")
    ap.add_argument("--reset", action="store_true", help="clear previous outputs before running")
    args = ap.parse_args()
    p = paths()
    if args.reset:
        reset_outputs(p)
    else:
        ensure_dirs(p)

    print(f"loading {DATA_CSV}")
    df = load_data(args.sample)
    print(f"rows: {len(df):,}")
    stage_baseline(df, p)
    print("checkpoint 1/10 baseline written")
    stage_winner_regions(df, p)
    print("checkpoint 2/10 winner regions written")
    df, pred_status = add_predictions(df)
    print(pred_status)
    stage_success(df, p)
    print("checkpoint 3/10 success indicators written")
    stage_failure(df, p)
    print("checkpoint 4/10 failure indicators written")
    stage_cases(df, p)
    print("checkpoint 5/10 representative cases written")
    stage_domain_patterns(df, p)
    print("checkpoint 6/10 domain patterns written")
    stage_interaction_patterns(df, p)
    print("checkpoint 7/10 interaction patterns written")
    stage_class_specific_patterns(df, p)
    print("checkpoint 8/10 class-specific patterns written")
    stage_worst_failure_anatomy(df, p)
    print("checkpoint 9/10 worst failure anatomy written")
    stage_synthesis(df, p, pred_status)
    print("checkpoint 10/10 synthesis written")
    print(f"markdown: {p.report_md}")
    print(f"docx: {p.report_docx}")


if __name__ == "__main__":
    main()
